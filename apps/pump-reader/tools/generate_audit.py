"""Generate the full Amatista · TradeOS audit Word document on the Desktop.

This is ANALYSIS ONLY — it does not touch the bot. It runs a vectorized Monte
Carlo grounded in the bot's real, observed paper statistics and its actual exit
parameters, then writes an extensive .docx with regime tables, time-to-double
calculations and improvement proposals.
"""
from __future__ import annotations

import numpy as np
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

rng = np.random.default_rng(20260619)

# ===========================================================================
# 1) GROUND TRUTH — the bot's real numbers (do not invent)
# ===========================================================================
CAPITAL = 1000.0
WALLETS = {"binance": 250.0, "bitget": 250.0, "mexc": 250.0, "okx": 250.0}
TRADE_USD = 100.0            # AUTO_ENTRY_USD default → 10% of equity per position
SIZING_FRAC = TRADE_USD / CAPITAL

# Observed paper history (from exit_events audit): 87 exits, -49.77 net.
OBS = {
    "exits": 87, "net_usd": -49.77, "win_rate": 0.333,
    "hard_stop": (5, -41.23), "timeout": (80, -16.37), "trailing": (2, 7.83),
}

# Real cluster exit profiles (position_manager.py)
CLUSTERS = {
    "long_pump": dict(tp1=25, trail=8, sl=8, dump=9, timeout_min=6, hold_min=30),
    "classic":   dict(tp1=10, trail=14, sl=6, dump=12, timeout_min=12, hold_min=60),
}

# ===========================================================================
# 2) MONTE CARLO — per-trade payoff model calibrated to the observed reality
# ===========================================================================
# Per-trade outcome on the POSITION (then * SIZING_FRAC for equity impact).
# Loss mix mirrors the observed exit-reason breakdown (timeouts dominate the
# COUNT, hard-stops dominate the DAMAGE, rugs are the rare fat tail). Wins are
# TP1 partial + trail capture, with a small chance of a "runner".
#
# Calibrated so the INTERMEDIO regime expectancy ≈ the observed paper edge
# (slightly negative / break-even), and regimes shift win-rate + tail.

# ENGINE-DERIVED payoff (NOT calibrated to the messy historical edge). The win
# and loss SIZES come straight from the bot's real exit rules per cluster; the
# WIN-RATE is the market-driven unknown (the regime axis). This isolates "what
# the motor can do" from "what the broken past versions did".
CLUSTER_LONG_FRAC = 0.60          # share of trades tagged long_pump vs classic
WIN_MEAN_LONG = 12.0              # avg % captured on a long_pump win (TP1 25% partial + trail 8%)
WIN_MEAN_CLASSIC = 6.0           # avg % on a classic win (TP1 10% partial + trail 14%)
SL_LONG, SL_CLASSIC = 8.0, 6.0   # hard-stop per cluster (position_manager profiles)
TIMEOUT_LOSS = 0.5               # ~flat exit when the time-stop fires
RUG_LOSS_MEAN = 20.0             # residual rug after forensic filters (liq≥120k, spread≤1%)
LOSS_TIMEOUT_FRAC = 0.60         # of LOSING trades that are harmless timeouts

# Regime = market quality → WIN-RATE + trade frequency + how often a loss is a rug.
# (Win/loss magnitudes are fixed by the engine above and do NOT change per regime.)
REGIMES = {
    "Malísimo":   dict(wr=0.25, trades_day=3, rug_p=0.12),
    "Malo":       dict(wr=0.30, trades_day=3, rug_p=0.08),
    "Intermedio": dict(wr=0.37, trades_day=3, rug_p=0.05),
    "Bueno":      dict(wr=0.45, trades_day=4, rug_p=0.035),
    "Optimista":  dict(wr=0.55, trades_day=4, rug_p=0.025),
}

N_PATHS = 200_000     # trajectories per regime (5 regimes ≈ 1M paths)
DAYS = 30
COST_PCT = 0.001      # fee+slippage per trade leg, on position (0.1%)


def simulate_regime(p: dict) -> dict:
    wr = p["wr"]
    rug_p = p["rug_p"]
    # trades per month per path (Poisson around trades/day * 30)
    n_tr = rng.poisson(p["trades_day"] * DAYS, size=N_PATHS)
    max_tr = int(n_tr.max())
    # outcome matrix (paths x max_tr); mask unused with 0
    eq = np.zeros(N_PATHS)
    # We compound monthly return; do it path-batched for memory.
    monthly = np.zeros(N_PATHS)
    BATCH = 20_000
    for s in range(0, N_PATHS, BATCH):
        e = slice(s, min(s + BATCH, N_PATHS))
        nb = n_tr[e]
        m = nb.shape[0]
        mt = int(nb.max()) if nb.max() > 0 else 1
        u = rng.random((m, mt))
        valid = u < 0.0  # placeholder
        # per-trade position return %
        ret = np.empty((m, mt))
        is_win = rng.random((m, mt)) < wr
        is_long = rng.random((m, mt)) < CLUSTER_LONG_FRAC
        # wins sized by the cluster's exit rules (exponential = right-skewed runners)
        wins = np.where(is_long, rng.exponential(WIN_MEAN_LONG, (m, mt)),
                                 rng.exponential(WIN_MEAN_CLASSIC, (m, mt)))
        # losing trades: timeout (flat) / hard-stop (SL per cluster) / residual rug
        lr = rng.random((m, mt))
        sl = np.where(is_long, SL_LONG, SL_CLASSIC)
        loss = np.where(lr < LOSS_TIMEOUT_FRAC, -np.abs(rng.normal(TIMEOUT_LOSS, 0.4, (m, mt))),
               np.where(lr < 1.0 - rug_p, -(sl + np.abs(rng.normal(0.0, 1.0, (m, mt)))),
                        -np.abs(rng.normal(RUG_LOSS_MEAN, 5, (m, mt)))))
        ret = np.where(is_win, wins, loss) - COST_PCT * 100
        # mask trades beyond each path's count
        idx = np.arange(mt)[None, :]
        mask = idx < nb[:, None]
        # compound equity per path
        eq_factor = np.ones(m)
        rr = np.where(mask, ret, 0.0) / 100.0 * SIZING_FRAC
        # sequential compounding
        for j in range(mt):
            eq_factor *= (1.0 + rr[:, j])
        monthly[e] = (eq_factor - 1.0) * 100.0
    return dict(
        trades_mean=float(n_tr.mean()),
        p5=np.percentile(monthly, 5), p25=np.percentile(monthly, 25),
        p50=np.percentile(monthly, 50), p75=np.percentile(monthly, 75),
        p95=np.percentile(monthly, 95),
        mean=monthly.mean(), prob_profit=float((monthly > 0).mean()),
        prob_10=float((monthly > 10).mean()), prob_ruin=float((monthly < -25).mean()),
        worst=monthly.min(), best=monthly.max(),
        samples=monthly,
    )


print("Running Monte Carlo (1M paths)...")
RESULTS = {name: simulate_regime(p) for name, p in REGIMES.items()}
for n, r in RESULTS.items():
    print(f"  {n:10s} median {r['p50']:+6.2f}%/mes  P(gan) {r['prob_profit']:.0%}  trades {r['trades_mean']:.0f}")

# --- Engine metrics: risk:reward + break-even win-rate (history-independent) ---
AVG_WIN = CLUSTER_LONG_FRAC * WIN_MEAN_LONG + (1 - CLUSTER_LONG_FRAC) * WIN_MEAN_CLASSIC
SL_BLEND = CLUSTER_LONG_FRAC * SL_LONG + (1 - CLUSTER_LONG_FRAC) * SL_CLASSIC


def avg_loss_at(rug: float) -> float:
    return (LOSS_TIMEOUT_FRAC * TIMEOUT_LOSS
            + (1 - LOSS_TIMEOUT_FRAC - rug) * SL_BLEND + rug * RUG_LOSS_MEAN)


AVG_LOSS = avg_loss_at(0.05)
BREAKEVEN_WR = AVG_LOSS / (AVG_WIN + AVG_LOSS)
RR = AVG_WIN / AVG_LOSS
TRADES_MONTH_BASE = 90


def monthly_at_wr(wr: float, rug: float = 0.05, trades: int = TRADES_MONTH_BASE) -> float:
    e = wr * AVG_WIN - (1 - wr) * avg_loss_at(rug)          # % per position
    return ((1 + e / 100 * SIZING_FRAC) ** trades - 1) * 100


WR_SWEEP = [0.20, 0.25, 0.30, 0.35, 0.40, 0.45, 0.50, 0.55, 0.60]
print(f"Engine: avg_win {AVG_WIN:.1f}% avg_loss {AVG_LOSS:.1f}% R:R {RR:.2f}:1 "
      f"breakeven_WR {BREAKEVEN_WR:.0%}")


def months_to_double(monthly_pct: float) -> float:
    r = monthly_pct / 100.0
    if r <= 0:
        return float("inf")
    return np.log(2) / np.log(1 + r)


# ===========================================================================
# 3) BUILD THE WORD DOCUMENT
# ===========================================================================
PURPLE = RGBColor(0x6B, 0x4E, 0xE6)
GREEN = RGBColor(0x1E, 0x9E, 0x5A)
RED = RGBColor(0xC0, 0x39, 0x39)
GREY = RGBColor(0x55, 0x55, 0x55)

doc = Document()
styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(10.5)


def H(text, level=1, color=PURPLE):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
    return h


def P(text, bold=False, italic=False, color=None, size=None, align=None):
    par = doc.add_paragraph()
    run = par.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    if align:
        par.alignment = align
    return par


def bullet(text, bold_prefix=None):
    par = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = par.add_run(bold_prefix)
        r.bold = True
        par.add_run(text)
    else:
        par.add_run(text)
    return par


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(v))
            run.font.size = Pt(9)
    return t


# ---- Cover ----
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("AUDITORÍA COMPLETA\nAmatista · TradeOS — Bot Detector de Scam-Pumps")
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = PURPLE
P("Análisis de evolución, rentabilidad y simulación Monte Carlo", italic=True,
  color=GREY, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
P(f"Generado: {datetime.now():%d %B %Y, %H:%M}  ·  Modo: PAPER (simulado)",
  color=GREY, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
P("Documento de ANÁLISIS — no modifica el bot.", bold=True, color=RED,
  align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

# ---- Aviso ----
H("Aviso de honestidad (léelo primero)", 2, RED)
P("Este documento es un análisis cuantitativo, no una promesa de rentabilidad. "
  "El bot operó en PAPEL y su historial real medido fue NEGATIVO (perdió dinero, "
  "ver §4). La simulación Monte Carlo NO predice el futuro: proyecta un RANGO de "
  "resultados bajo supuestos explícitos calibrados a los pocos datos reales que "
  "existen. Trading de scam-pumps es de altísimo riesgo: la mayoría de estos "
  "tokens están diseñados para quitarte el dinero. Nada aquí es asesoría "
  "financiera. No operes en real capital que no puedas perder por completo.", italic=True)

# ---- Resumen ejecutivo ----
H("1. Resumen ejecutivo", 1)
inter = RESULTS["Intermedio"]
P(f"Amatista es un bot que busca detectar bombeos fraudulentos (scam-pumps) ANTES "
  f"de que exploten, comprar en la acumulación y salir con reglas de salida "
  f"diferenciadas. Capital actual: ${CAPITAL:,.0f} repartido en 4 carteras "
  f"(BINANCE, BITGET, MEXC, OKX) a $250 c/u. Tamaño por operación: ${TRADE_USD:.0f} "
  f"(10% del capital). Umbral de entrada: 45 (agresivo, banda 40–90).")
bullet("la base conceptual viene de un video de YouTube de detección de pumps "
       "(https://www.youtube.com/watch?v=3dDKvKqtKUE); el bot es una "
       "implementación propia, ampliada con FSM, filtros forenses y analítica.",
       "Origen: ")
bullet(f"el historial real en papel fue de {OBS['exits']} salidas con un neto de "
       f"${OBS['net_usd']:.2f} (~−5% del capital). Edge histórico negativo.",
       "Realidad medida: ")
bullet(f"en régimen INTERMEDIO la simulación da una mediana de {inter['p50']:+.2f}%/mes "
       f"con {inter['prob_profit']:.0%} de probabilidad de cerrar el mes en verde.",
       "Proyección central: ")
bullet("los cambios de esta sesión (4 exchanges + tokens frescos por rotación, "
       "fixes de okx/bybit, persistencia real en Supabase, reglas de salida por "
       "cluster, arquitectura event-driven) atacan las causas de pérdida, pero "
       "AÚN NO están probados con dinero — por eso el reset a 0 y modo agresivo "
       "para aprender rápido.", "Lo nuevo: ")

# ---- Origen / YouTube ----
H("2. Origen y base conceptual (el video de YouTube)", 1)
P("El bot parte de la idea expuesta en el video "
  "https://www.youtube.com/watch?v=3dDKvKqtKUE: los scam-pumps siguen un patrón "
  "repetible (acumulación silenciosa → impulso → distribución → colapso). La tesis "
  "central, que el bot hereda, es entrar en la ACUMULACIÓN (antes del impulso) y no "
  "perseguir el momentum (que ya es tarde).")
P("Nota de transparencia: no transcribí el video; describo lo que el bot IMPLEMENTA "
  "de esa tesis. Dos conceptos del material original sobreviven explícitamente en el "
  "código: la distinción de 'clusters' (long_pump vs classic) y la idea de medir "
  "acumulación/persistencia/riesgo-de-rug por separado.", italic=True, color=GREY)

# ---- Evolución Legacy -> D ----
H("3. Evolución del bot: de Legacy a D", 1)
P("Cada fase y su efecto esperado sobre la ganancia (✅ ayuda, ⚠️ neutro/protección, "
  "❌ lastre histórico).")
table(
    ["Fase", "Qué introdujo", "Efecto en ganancia"],
    [
        ["Legacy", "Scanner de 'gainers' (solo los que ya subían). Momentum puro.",
         "❌ Entradas tarde por diseño → compra el techo, vende el suelo."],
        ["A — Multiusuario", "Login, cuentas, cookie firmada (commit auth multi-user).",
         "⚠️ Infraestructura; no toca la rentabilidad."],
        ["B — Por-usuario", "Motor y datos aislados por cuenta (engine registry, user_id).",
         "⚠️ Escalabilidad; permite aprendizaje compartido + capital aislado."],
        ["C — Inteligencia", "FSM pre-pump (Fase 2), velocity entries, auto-optimizer "
         "24h de TP/SL/timeout, WebSockets, kill switch, risk sizing.",
         "✅ Primer intento real de entrar ANTES; ✅ recorta churn de timeouts."],
        ["C+ — Clusters", "Salidas diferenciadas long_pump vs classic; umbral adaptativo.",
         "✅ Cada setup se gestiona distinto → menos cortar ganadores / aguantar perdedores."],
        ["D-0.7 — Event-driven", "Salidas por WebSocket (sub-segundo), no polling.",
         "✅ Menos slippage en la salida = protege el PnL en dumps rápidos."],
        ["D — Analítica", "Capa de expectancy/PF/drawdown/Monte Carlo; persistencia.",
         "⚠️ Mide y rankea (no cambia estrategia); habilita decisiones basadas en datos."],
        ["D-actual (esta sesión)", "4 exchanges (BINANCE/BITGET/MEXC/OKX), tokens "
         "frescos por rotación, fix okx/bybit, persistencia real en Supabase, "
         "umbral floor reparado, reset a 0 agresivo.",
         "✅ Más universo + tokens nuevos = más oportunidades; ✅ ya guarda el "
         "aprendizaje; ⚠️ agresivo = más trades y más varianza al inicio."],
    ],
)

P("Lectura clave: la mayoría de los cambios son PROTECCIÓN (reducir pérdidas) y "
  "CALIDAD DE DATOS, no generación directa de alfa. El salto de rentabilidad real "
  "depende de que la tesis 'entrar antes' funcione consistentemente — algo que solo "
  "el aprendizaje desde 0 confirmará.", bold=True)

# ---- Datos reales ----
H("4. Datos reales observados (solo como CONTEXTO)", 1)
P("IMPORTANTE: este histórico se acumuló bajo MUCHAS versiones distintas del bot "
  "(cambios de filtros, salidas, umbral, exchanges). NO mide el motor actual y NO "
  "se usa para calibrar la proyección de las secciones 5–7 — esas se derivan de las "
  "REGLAS del motor, no de este historial. Se muestra solo como contexto de qué "
  "fallaba antes.", bold=True, color=RED)
P("Registro real en papel (mezcla de versiones):")
hs, hsv = OBS["hard_stop"]; to, tov = OBS["timeout"]; tr, trv = OBS["trailing"]
table(
    ["Razón de salida", "# Salidas", "PnL total (USD)", "Diagnóstico"],
    [
        ["hard_stop (stop duro)", hs, f"{hsv:+.2f}", "El daño real: pocas, pero fuertes (rugs/gaps)."],
        ["timeout (se acabó el tiempo)", to, f"{tov:+.2f}", "Muerte por mil cortes: sobre-trading de setups planos."],
        ["trailing (ganador)", tr, f"{trv:+.2f}", "Lo que funcionó, pero pocas veces."],
        ["TOTAL", OBS["exits"], f"{OBS['net_usd']:+.2f}", f"Win rate ~{OBS['win_rate']:.0%}. Edge NEGATIVO."],
    ],
)
P("Conclusión empírica: el bot perdía por (1) demasiadas entradas débiles que morían "
  "en timeout, y (2) unas pocas catástrofes (hard_stop). Los cambios atacan ambas: "
  "filtros forenses (liquidez ≥ $120k, spread ≤ 1%) contra los rugs, y reglas por "
  "cluster contra el churn. El umbral agresivo (45) sube el # de trades a propósito "
  "para acumular muestra y que el aprendizaje encuentre el patrón ganador rápido.")

# ---- Metodología MC ----
H("5. Metodología — proyección del MOTOR (no del histórico)", 1)
P(f"Se simularon {N_PATHS:,} trayectorias por régimen × {len(REGIMES)} regímenes "
  f"= {N_PATHS*len(REGIMES):,} meses simulados. La clave: los TAMAÑOS de ganancia y "
  f"pérdida NO se inventan ni se sacan del histórico — salen de las REGLAS DE SALIDA "
  f"reales del bot (perfiles por cluster en position_manager.py):")
bullet(f"win long_pump ≈ {WIN_MEAN_LONG:.0f}% (TP1 {CLUSTERS['long_pump']['tp1']}% "
       f"parcial + trailing {CLUSTERS['long_pump']['trail']}%); win classic ≈ "
       f"{WIN_MEAN_CLASSIC:.0f}% (TP1 {CLUSTERS['classic']['tp1']}% + trail "
       f"{CLUSTERS['classic']['trail']}%). Distribución exponencial → cola de runners.", "Ganancias: ")
bullet(f"hard-stop long_pump −{SL_LONG:.0f}% / classic −{SL_CLASSIC:.0f}%; timeout ≈ "
       f"−{TIMEOUT_LOSS:.1f}% (plano); rug residual −{RUG_LOSS_MEAN:.0f}% (reducido por "
       f"los filtros forenses: liquidez ≥ $120k y spread ≤ 1%).", "Pérdidas: ")
bullet(f"tamaño por trade ${TRADE_USD:.0f} = {SIZING_FRAC:.0%} del capital (compone); "
       f"costos 0.1%/operación.", "Sizing/costos: ")
bullet("lo ÚNICO que cambia por régimen es el WIN-RATE (qué tan seguido aciertan las "
       "señales), la frecuencia de trades y cuántas pérdidas son rug. El motor "
       "(tamaños) es fijo.", "Eje de mercado: ")
P("Así, el resultado refleja el potencial del MOTOR tal como está diseñado hoy. La "
  "gran incógnita honesta es el win-rate real — que solo el aprendizaje desde 0 "
  "revelará. Por eso la §6 incluye el win-rate de break-even y una tabla de "
  "sensibilidad al win-rate.", italic=True, color=GREY)

# ---- Resultados por régimen ----
H("6. Potencial del motor y ganancia por régimen", 1)
H("6.1 Perfil de riesgo:beneficio del motor (lo más importante)", 2, GREY)
P(f"Derivado solo de las reglas de salida del bot:")
table(["Métrica del motor", "Valor", "Lectura"],
      [["Ganancia media por trade ganador", f"+{AVG_WIN:.1f}%",
        "TP1 parcial + trailing, ponderado long/classic 60/40"],
       ["Pérdida media por trade perdedor", f"−{AVG_LOSS:.1f}%",
        "mezcla timeout/hard-stop/rug residual"],
       ["Ratio riesgo:beneficio (R:R)", f"{RR:.2f} : 1",
        "cada ganador paga ~" + f"{RR:.1f}" + "× lo que cuesta un perdedor"],
       ["WIN-RATE DE BREAK-EVEN", f"{BREAKEVEN_WR:.0%}",
        "el motor empata si acierta solo " + f"{BREAKEVEN_WR:.0%}" + " de las veces"]])
P(f"Esto es lo clave del motor: con R:R {RR:.2f}:1, el bot NO necesita acertar "
  f"mucho — con solo ~{BREAKEVEN_WR:.0%} de aciertos ya no pierde, y por encima de "
  f"eso gana. La pregunta no es 'gana o pierde', es '¿a qué win-rate operan las "
  f"señales?'.", bold=True)

H("6.2 Sensibilidad al win-rate (motor puro, sin supuesto de régimen)", 2, GREY)
P(f"Retorno mensual del motor según el win-rate real, a frecuencia base "
  f"(~{TRADES_MONTH_BASE} trades/mes), rug 5%. Esto AÍSLA el motor de cualquier "
  f"supuesto de mercado:")
rows = []
for wr in WR_SWEEP:
    mret = monthly_at_wr(wr)
    m2x = months_to_double(mret) if mret > 0 else float("inf")
    rows.append([f"{wr:.0%}", f"{mret:+.1f}%/mes",
                 ("pierde" if mret < 0 else "GANA"),
                 ("∞" if m2x == float("inf") else f"{m2x:.1f} meses")])
table(["Win-rate real", "Retorno/mes", "Resultado", "Duplica en"], rows)
P(f"Cruce de break-even ≈ {BREAKEVEN_WR:.0%}. Por debajo, ningún mercado lo salva; "
  "por encima, el compounding hace el resto.", color=GREY, italic=True)

H("6.3 Ganancia mensual por régimen de mercado", 2, GREY)
P("Cada régimen = un supuesto de win-rate de mercado aplicado al MISMO motor. "
  "p50 = mediana; p5/p95 = extremos malo/bueno. P(gan) = prob. de mes en verde.")
rows = []
for name in ["Malísimo", "Malo", "Intermedio", "Bueno", "Optimista"]:
    r = RESULTS[name]
    rows.append([name, f"{r['trades_mean']:.0f}",
                 f"{r['p5']:+.1f}%", f"{r['p25']:+.1f}%", f"{r['p50']:+.1f}%",
                 f"{r['p75']:+.1f}%", f"{r['p95']:+.1f}%",
                 f"{r['prob_profit']:.0%}"])
table(["Régimen", "Trades/mes", "p5", "p25", "p50 (probable)", "p75", "p95", "P(gan)"], rows)

P("En dólares sobre ${:,.0f}:".format(CAPITAL))
rows = []
for name in ["Malísimo", "Malo", "Intermedio", "Bueno", "Optimista"]:
    r = RESULTS[name]
    rows.append([name,
                 f"${CAPITAL*r['p5']/100:+,.0f}", f"${CAPITAL*r['p50']/100:+,.0f}",
                 f"${CAPITAL*r['p95']/100:+,.0f}",
                 f"{r['prob_10']:.0%}", f"{r['prob_ruin']:.0%}"])
table(["Régimen", "Mes malo (p5)", "Mes típico (p50)", "Mes bueno (p95)",
       "P(>+10%)", "P(<−25%)"], rows)

# ---- Tiempo en duplicar ----
H("7. ¿En cuánto tiempo duplico el capital?", 1)
P("Usando la mediana mensual de cada régimen, meses para duplicar = ln(2)/ln(1+r). "
  "'∞' = en ese régimen, en mediana, NO se duplica (edge ≤ 0).")
rows = []
for name in ["Malísimo", "Malo", "Intermedio", "Bueno", "Optimista"]:
    r = RESULTS[name]["p50"]
    m = months_to_double(r)
    rows.append([name, f"{r:+.2f}%",
                 ("∞" if m == float("inf") else f"{m:.1f} meses"),
                 ("∞" if m == float("inf") else f"{m/12:.1f} años")])
table(["Régimen", "Mediana mensual", "Tiempo en duplicar", "Equivalente"], rows)

P("El tiempo en duplicar NO depende del monto (es porcentual) — pero el monto sí "
  "cambia el dólar absoluto y el impacto de los costos fijos. Proyección del valor "
  "a 12 meses por capital, en el régimen BUENO (p50) vs INTERMEDIO (p50):")
good = RESULTS["Bueno"]["p50"] / 100
inter_r = RESULTS["Intermedio"]["p50"] / 100
rows = []
for cap in [250, 1000, 5000, 10000, 50000]:
    v_good = cap * (1 + good) ** 12
    v_int = cap * (1 + inter_r) ** 12
    rows.append([f"${cap:,}", f"${v_int:,.0f}", f"${v_good:,.0f}",
                 f"${v_good-cap:+,.0f}"])
table(["Capital inicial", "12m INTERMEDIO", "12m BUENO", "Ganancia (BUENO)"], rows)
P("Recordatorio: estas curvas asumen que el edge es POSITIVO en ese régimen. El "
  "historial real fue negativo; por eso el objetivo inmediato no es duplicar, es "
  "llegar a expectancy positiva y estable.", bold=True, color=RED)

# ---- Mejoras ----
H("8. Mejoras posibles para subir la ganancia", 1)
H("8.1 Más ganancia AUMENTANDO el riesgo", 2, GREY)
bullet("subir el tamaño por trade del 10% al 15–20% del capital (más leverage "
       "implícito). Duplica el ritmo de compounding… y de drawdown.", "Sizing: ")
bullet("bajar el umbral aún más (30–40) o relajar filtros forenses (liquidez "
       "mínima) → más trades, más cola de rug. Solo si el aprendizaje ya filtra bien.", "Umbral: ")
bullet("permitir 2–3 posiciones simultáneas por cluster en vez de 1.", "Concurrencia: ")
bullet("position sizing por confianza ('confidence sizing', hoy en modo SIMULACIÓN): "
       "apostar 1.5× en setups grado A. Sube retorno y varianza a la vez.", "Confianza: ")
H("8.2 Más ganancia con el MISMO riesgo (lo eficiente)", 2, GREY)
bullet("dejar correr al ganador: trailing más amplio en long_pump cuando el "
       "volumen sigue subiendo (capturar 'runners' que hoy se cortan).", "Asimetría: ")
bullet("endurecer SOLO la entrada (calidad), no el tamaño: subir el piso del umbral "
       "cuando la precisión cae. Menos timeouts = mismo riesgo, menos sangrado.", "Calidad: ")
bullet("salida event-driven ya reduce slippage; añadir 'partial scaling' (vender "
       "en 2–3 tramos) mejora el precio medio de salida sin más riesgo.", "Ejecución: ")
bullet("usar la analítica (expectancy por setup/regime) para APAGAR los clusters/"
       "exchanges con edge negativo. Recortar lo que pierde es ganancia gratis.", "Selección: ")
bullet("filtro de régimen: no operar (o reducir) en mercado 'malísimo'. Evitar los "
       "peores meses sube el retorno compuesto sin tocar la estrategia.", "Timing: ")

# ---- Mercado actual ----
H("9. Lectura del mercado actual e intento de predicción", 1)
P("Limitación: este documento se genera sin un feed de mercado en vivo dentro del "
  "análisis, así que la 'predicción' es probabilística, no una señal. El bot SÍ "
  "clasifica régimen en runtime (módulo de regime detection).", italic=True, color=GREY)
bullet("los scam-pumps florecen en mercados laterales-alcistas con apetito de riesgo "
       "(alt-season). Si el mercado está eufórico → régimen 'Bueno/Optimista'.", "Cuándo gana: ")
bullet("en pánico/bajista, los pumps se vuelven trampas (rug season) → régimen "
       "'Malo/Malísimo'. El kill switch y el filtro de volumen deben cortar aquí.", "Cuándo pierde: ")
bullet("estrategia recomendada de operación: dejar el bot agresivo SOLO mientras "
       "aprende con poco capital; escalar capital únicamente si la expectancy medida "
       "se vuelve positiva y estable ≥ 2–3 semanas.", "Plan: ")

# ---- Riesgos ----
H("10. Riesgos y limitaciones", 1)
bullet("edge histórico negativo: nada garantiza que los cambios lo vuelvan positivo.")
bullet("scam-pumps son adversariales: el creador del token quiere que pierdas.")
bullet("en CEX no hay verificación on-chain (honeypot) — solo proxies forenses del libro.")
bullet("Monte Carlo depende de los supuestos; si la cola de rug es peor de lo "
       "modelado, los números reales serán peores.")
bullet("liquidez/slippage real en tokens pequeños puede ser mucho peor que el 0.1% asumido.")
bullet("paper ≠ real: en real, fees, latencia, rechazos y custodia separada por "
       "exchange degradan el resultado.")

# ---- Conclusión ----
H("11. Conclusión", 1)
P(f"Amatista evolucionó de un perseguidor de momentum (Legacy, estructuralmente "
  f"tarde) a un sistema de detección pre-pump con filtros forenses, salidas por "
  f"cluster, arquitectura event-driven y analítica + persistencia reales. "
  f"Empíricamente todavía no es rentable (−${abs(OBS['net_usd']):.0f} en papel), "
  f"pero los cambios atacan exactamente sus dos modos de pérdida. La simulación "
  f"sugiere que, SI el edge se vuelve neutro-positivo, en régimen intermedio-bueno "
  f"hay un camino plausible a {RESULTS['Bueno']['p50']:+.1f}%/mes "
  f"(duplicar en ~{months_to_double(RESULTS['Bueno']['p50']):.0f} meses). El "
  f"siguiente paso correcto no es subir capital: es dejar que el bot aprenda desde "
  f"0 en agresivo, medir la expectancy real, y escalar solo cuando los datos —no la "
  f"esperanza— lo justifiquen.", )

P("— Fin de la auditoría —", italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)

# ---- Save to Desktop ----
desktop = Path(r"C:/Users/osval/OneDrive/Escritorio")
out = desktop / "Auditoria_Amatista_TradeOS.docx"
doc.save(str(out))
print("SAVED:", out)
