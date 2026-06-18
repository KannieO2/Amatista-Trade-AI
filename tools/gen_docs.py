# -*- coding: utf-8 -*-
"""Genera la documentacion tecnica completa de TradeOS Amatista en .docx.
Todo el contenido sale del codigo real (no inventado). Honestidad sobre fuentes
de datos: lo que NO esta cableado se marca como tal."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INK = RGBColor(0x0c, 0x10, 0x18)
PINK = RGBColor(0xff, 0x2f, 0x6e)
GREY = RGBColor(0x55, 0x5b, 0x66)
GREEN = RGBColor(0x18, 0x8a, 0x5a)
RED = RGBColor(0xc0, 0x2a, 0x2a)

doc = Document()

# Base style
base = doc.styles['Normal']
base.font.name = 'Calibri'
base.font.size = Pt(10.5)

def _shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)

def h1(t):
    p = doc.add_heading(t, level=1)
    for r in p.runs: r.font.color.rgb = PINK
    return p

def h2(t):
    p = doc.add_heading(t, level=2)
    for r in p.runs: r.font.color.rgb = INK
    return p

def h3(t):
    p = doc.add_heading(t, level=3)
    for r in p.runs: r.font.color.rgb = GREY
    return p

def p(t, bold=False, italic=False, color=None, size=10.5):
    par = doc.add_paragraph()
    r = par.add_run(t); r.bold = bold; r.italic = italic; r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    return par

def bullet(t, bold_prefix=None):
    par = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = par.add_run(bold_prefix); r.bold = True
        par.add_run(t)
    else:
        par.add_run(t)
    return par

def code(t):
    par = doc.add_paragraph()
    par.paragraph_format.left_indent = Inches(0.2)
    par.paragraph_format.space_before = Pt(2); par.paragraph_format.space_after = Pt(6)
    for line in t.split('\n'):
        r = par.add_run(line + '\n')
        r.font.name = 'Consolas'; r.font.size = Pt(8.8)
        r.font.color.rgb = RGBColor(0x1b, 0x2a, 0x4a)
    return par

def table(headers, rows, widths=None):
    tb = doc.add_table(rows=1, cols=len(headers))
    tb.style = 'Light Grid Accent 1'
    tb.alignment = WD_TABLE_ALIGNMENT.LEFT
    hc = tb.rows[0].cells
    for i, htext in enumerate(headers):
        hc[i].text = ''
        rr = hc[i].paragraphs[0].add_run(htext); rr.bold = True; rr.font.size = Pt(9)
        rr.font.color.rgb = RGBColor(0xff,0xff,0xff)
        _shade(hc[i], '2f3a52')
    for row in rows:
        cells = tb.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            rr = cells[i].paragraphs[0].add_run(str(val)); rr.font.size = Pt(8.8)
    if widths:
        for i, w in enumerate(widths):
            for row in tb.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return tb

# ============================ PORTADA ============================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('Amatista · TradeOS'); r.bold = True; r.font.size = Pt(30); r.font.color.rgb = PINK
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Documentación Técnica Completa del Bot'); r.font.size = Pt(15); r.font.color.rgb = INK
sub2 = doc.add_paragraph(); sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run('Detector de pumps fraudulentos (scam-pump) + motor de salida gestionada'); r.italic = True; r.font.size = Pt(11); r.font.color.rgb = GREY
meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run('Generado: 2026-06-18  ·  Modo por defecto: PAPER  ·  apps/pump-reader'); r.font.size = Pt(9); r.font.color.rgb = GREY

p('')
p('Nota de honestidad (regla del proyecto): este documento describe EXACTAMENTE lo que '
  'el código hace hoy. Cualquier fuente de datos que NO está cableada (CoinMarketCap, '
  'DexScreener, on-chain, holders, social) se marca explícitamente como NO integrada. '
  'No se inventan datos de negocio.', italic=True, color=GREY, size=9.5)

# Veredicto destacado
p('')
vb = doc.add_paragraph()
vb.paragraph_format.left_indent = Inches(0.15)
rr = vb.add_run('VEREDICTO CUANTITATIVO (Monte Carlo 1.000.000 casos): bajo supuestos realistas la '
                'estrategia es de expectativa NEGATIVA (EV ≈ -1,23%/trade). El valor real del bot es la '
                'DETECCIÓN y el FILTRADO (los gates aportan +4,27%/trade), no el auto-trading. '
                'Mantener en PAPER. Detalle en la sección 19.')
rr.bold = True; rr.font.size = Pt(10); rr.font.color.rgb = RED

doc.add_page_break()

# ============================ INDICE ============================
h1('Contenido')
idx = [
 '1. Arquitectura general y componentes',
 '2. Flujo completo: de la señal al cierre',
 '3. Datos que consume (APIs, exchanges, CoinGecko, CMC, DexScreener, on-chain, holders, social)',
 '4. Reglas de filtrado (universo de tokens)',
 '5. Cómo se calcula el score (pump_score, confidence, clasificación, cluster)',
 '6. Reglas de ENTRADA exactas',
 '7. Reglas de SALIDA exactas (motor PositionManager)',
 '8. Stop loss (hard stop)',
 '9. Trailing stop',
 '10. Take profits',
 '11. Dump detector y time-stop dinámico',
 '12. Gestión de capital',
 '13. Gestión de riesgo (RiskGuard)',
 '14. Cooldowns',
 '15. Filtros anti-rug (ForensicFilter)',
 '16. Sistema de aprendizaje',
 '17. Todas las fórmulas',
 '18. Todos los parámetros configurables (variables de entorno)',
 '19. Ejemplo real: una señal de principio a fin + resultados Monte Carlo',
 '20. Por qué existe cada regla',
]
for line in idx:
    p(line)
doc.add_page_break()

# ============================ 1 ============================
h1('1. Arquitectura general y componentes')
p('TradeOS Amatista es UNA sola app FastAPI (apps/pump-reader) con dashboard HTML embebido. '
  'Es multi-tenant: cada cuenta es su PROPIO bot independiente (balance, posiciones, riesgo, '
  'equity y P&L propios). Lo único compartido es el "cerebro": el escáner de mercado y el '
  'aprendizaje del pump.')
h3('Módulos (app/)')
table(['Módulo', 'Responsabilidad'], [
 ['main.py', 'API FastAPI, loops de fondo (scan/monitor/velocity), orquestación, endpoints'],
 ['scanner.py', 'Escaneo de mercado real vía CCXT, scoring por reglas, ForensicFilter'],
 ['velocity.py', 'Disparador en tiempo real por aceleración de volumen (entre escaneos)'],
 ['position_manager.py', 'Motor de SALIDA: TP1, trailing, dump, hard-stop, break-even, time-stop'],
 ['executor.py', 'Capa de ejecución: PaperBroker (simulado) / LiveBroker (CCXT real)'],
 ['risk.py', 'RiskGuard: ningún orden pasa sin aprobarse (caps + kill switch)'],
 ['user_bot.py', 'Registro por usuario: un UserBot por cuenta (estado aislado)'],
 ['account.py', 'Balance real READ-ONLY vía CCXT fetch_balance (claves del usuario)'],
 ['market.py', 'CoinGecko: FDV / market cap / supply (solo bajo demanda)'],
 ['learning.py', 'LearningLab: MFE/MAE, lead-time, precisión/recall, propuestas de umbral'],
 ['scanner→store', 'Persistencia Supabase (service key, RLS); el navegador nunca habla con la DB'],
], widths=[1.6, 4.6])

h3('Loops de fondo (asyncio, arrancan en lifespan)')
table(['Loop', 'Cadencia (def.)', 'Qué hace'], [
 ['_auto_scan_loop', '300 s', 'Escanea mercado, genera candidatos, dispara alertas/auto-entradas'],
 ['_velocity_loop', '10 s', 'Vigila hot-list, entra al acelerar el volumen (no espera al escaneo)'],
 ['_monitor_loop', '15 s', 'Tickea cada posición abierta de cada bot → ejecuta salidas'],
 ['_account_loop', '120 s', 'Refresca balance real (solo si el owner tiene claves)'],
 ['_grid_tick_loop', '15 s', 'Avanza el grid bot GRVT (producto separado, mismo app)'],
 ['_grid_sync_loop', '60 s', 'Espeja estado del grid a Supabase'],
 ['_daily_discover_loop', '86400 s', 'Escaneo diario + log fechado'],
], widths=[1.6, 1.1, 3.5])

# ============================ 2 ============================
h1('2. Flujo completo: de la señal al cierre')
p('Hay dos caminos de entrada (scan lento + velocity rápido) que convergen en el mismo '
  'embudo de gates y, una vez dentro, en el mismo motor de salida.')

h3('Fase A — Descubrimiento (cada 300 s, _perform_scan)')
steps_a = [
 'scan_markets() corre en paralelo sobre los exchanges configurados (binance, mexc, bitget).',
 'Por exchange: load_markets() + fetch_tickers() (API pública CCXT, sin claves).',
 'Filtro de universo: solo spot activo, quote USDT, NO majors/stables/leveraged, volumen 24h en [100k, 60M] USD, cambio 24h > 0.',
 'Se ordenan los gainers por % desc y se toma el shortlist (20 tokens).',
 'Deep-scan de cada uno (concurrencia 5): fetch_ohlcv 1h (24 velas) + fetch_order_book (50 niveles).',
 'Se calculan: volume_spike, imbalance, liquidity_usd, spread_pct, top_book_share.',
 'score_candidate() → pump_score, confidence_score, classification, cluster, flags.',
 'forensic_check() marca (no falsea) manipulación: spread/liquidez/concentración del libro.',
 'Solo se conservan candidatos con pump_score ≥ 1.',
]
for s in steps_a: bullet(s)

h3('Fase B — Confirmación y alerta')
steps_b = [
 'status = waiting_confirmation si pump_score ≥ umbral adaptativo (arranca en 75).',
 'Si se confirma: alerta Telegram + persistencia (alert, pump_candidate, bot_log) + record_alert en LearningLab.',
 'En modo paper: cada UserBot con auto_entry activo y que NO tenga ya el símbolo → _auto_enter().',
]
for s in steps_b: bullet(s)

h3('Fase C — Entrada (gates de _auto_enter, en orden)')
steps_c = [
 '(1) Confianza: confidence_score ≥ 50, si no → skip_low_confidence.',
 '(2) Anti-chase: price_change_pct_24h < 60%, si no → skip_exhausted.',
 '(3) Volumen (solo camino scan): volume_spike ≥ 2,5x, si no → skip_low_volume. El camino velocity lo omite (ya probó aceleración en vivo).',
 '(4) ForensicFilter: spread ≤ 1,0% AND liquidez ≥ 120.000 USD AND top-3 del libro ≤ 80%. Si bloquea → forensic_block.',
 'Si pasa: engine.act() → RiskGuard.evaluate() → PaperBroker.place() (slippage 0,5%) → pm.open().',
]
for s in steps_c: bullet(s)

h3('Camino rápido — Velocity (cada 10 s, paralelo)')
p('accel = volumen_1m_última_vela_cerrada / media(volúmenes_1m_previos). Dispara si '
  'accel ≥ 4 AND precio subiendo AND vela nueva AND símbolo fuera de cooldown (600 s). '
  'Al disparar entra por _auto_enter con accel, saltándose el piso de volumen.')

h3('Fase D — Gestión y salida (cada 15 s, _monitor_loop → pm.step)')
p('Por cada posición abierta se obtiene precio en vivo (fetch_price) y volumen 1m '
  '(fetch_1m_volume) y se evalúan las salidas en este ORDEN de prioridad:')
table(['Prioridad', 'Salida', 'Condición', 'Acción'], [
 ['1', 'hard_stop', 'gain ≤ -8%', 'vende 100%'],
 ['2', 'dump', 'caída de 1 tick ≥ 10%', 'vende 100%'],
 ['3', 'break_even', 'armado a +4%; cae a entry+0,5%', 'vende 100%'],
 ['4', 'timeout', 'lateral (|gain|≤3%) + volumen agotado', 'vende 100%'],
 ['5', 'tp1', 'gain ≥ +15% (fase 1)', 'vende 60%, pasa a fase 2'],
 ['6', 'trailing', 'caída desde el pico ≥ 10% (fase 2)', 'vende el resto'],
], widths=[0.8, 1.1, 2.6, 1.4])
p('Al cierre total: se grada la calidad de entrada (early/perfect/late) y se ajusta el '
  'umbral adaptativo (loss-averse). Toda salida se persiste y notifica.')

# ============================ 3 ============================
h1('3. Datos que consume')
p('Principio del código (scanner.py): "No machine learning, no invented data. Any signal '
  'that cannot be sourced from the public exchange is omitted, not faked." Por eso varias '
  'de las fuentes que pediste NO están cableadas — y se dice abiertamente.')

table(['Fuente', 'Estado', 'Qué se usa / por qué NO'], [
 ['CCXT (exchanges)', 'ACTIVA (núcleo)', 'API pública sin claves: fetch_tickers, fetch_ohlcv (1h scan / 1m velocity+exit), fetch_order_book (50/20 niveles). Es la fuente de TODAS las señales.'],
 ['Binance', 'ACTIVA', 'Exchange escaneado (lista pocos microcaps).'],
 ['MEXC', 'ACTIVA', 'Exchange escaneado + de ejecución por defecto (aquí viven los scam-pumps).'],
 ['Bitget', 'ACTIVA', 'Exchange escaneado + de ejecución por defecto.'],
 ['CoinGecko', 'ACTIVA (limitada)', 'API free /coins/markets: FDV, market cap, supply. SOLO bajo demanda (modal del token, endpoint /token/market). NO entra en el score ni en la decisión de entrada. Caché 30 min.'],
 ['CoinMarketCap', 'NO integrada', 'No hay ninguna llamada a CMC en el código. Si se quisiera, sería redundante con CoinGecko.'],
 ['DexScreener', 'NO integrada', 'Comentario explícito en scanner.py: hook opcional dejado a propósito, no adivinado. En un CEX se opera contra el libro del exchange, no contra un contrato.'],
 ['On-chain', 'NO disponible', 'En un CEX no hay contrato/wallets visibles. Se omite, no se falsea. Proxy honesto: profundidad del orderbook.'],
 ['Holders', 'NO disponible', 'Igual que on-chain: no hay lista de holders en un CEX. Proxy: concentración top-3 del libro (top_book_share).'],
 ['Social', 'NO integrada', 'No hay scraping de Twitter/Telegram/etc. en el código.'],
 ['CCXT fetch_balance', 'ACTIVA (read-only)', 'Balance real del usuario con SUS claves spot (sin permiso de retiro). Solo lee; nunca opera ni retira.'],
 ['Supabase', 'ACTIVA', 'Persistencia (service key + RLS). Candidatos, posiciones, salidas, equity, logs, aprendizaje.'],
 ['Telegram', 'ACTIVA', 'Notificaciones (alertas, entradas, salidas, errores).'],
], widths=[1.4, 1.2, 3.7])

p('Resumen honesto: el bot es un detector de patrones de CEX. Las "checks anti-rug de '
  'DEX" (honeypot, liquidez quemada, impuesto compra/venta, lista de holders) NO tienen '
  'fuente aquí y NO se inventan. Lo medible del libro público del CEX sí se aplica.',
  italic=True, color=GREY)

# ============================ 4 ============================
h1('4. Reglas de filtrado (universo de tokens)')
h3('Exclusiones de base (_is_altcoin)')
bullet('Debe ser spot y activo; quote == USDT.', None)
bullet('NO majors: BTC, ETH, BNB, SOL, XRP, ADA, DOGE, TRX, DOT, AVAX, LINK, MATIC, LTC, BCH, ATOM, ETC, XLM. (demasiado líquidos para "pumpearse")', None)
bullet('NO stables/fiat: USDC, FDUSD, TUSD, DAI, USDP, BUSD, USD1, AEUR, EUR, TRY, BRL, ARS, GBP, JPY, EURI, XUSD.', None)
bullet('NO tokens apalancados: marcadores UP, DOWN, BULL, BEAR, 3L, 3S, 5L, 5S (son derivados, no pumps spot).', None)
h3('Cotas de shortlist')
table(['Filtro', 'Valor', 'Razón'], [
 ['MIN_QUOTE_VOLUME_USD', '100.000', 'Ignora polvo (no operable)'],
 ['MAX_QUOTE_VOLUME_USD', '60.000.000', 'Ignora blue chips de mega-volumen'],
 ['cambio 24h', '> 0', 'Solo gainers (un pump sube)'],
 ['SHORTLIST_SIZE', '20', 'Top gainers a deep-scan por exchange'],
 ['DEEP_FETCH_CONCURRENCY', '5', 'Llamadas OHLCV+libro en paralelo'],
 ['pump_score', '≥ 1', 'Se descarta lo que no puntúa'],
], widths=[2.2, 1.3, 2.8])

# ============================ 5 ============================
h1('5. Cómo se calcula el score')
p('Función score_candidate() en scanner.py. Calcula DOS criterios competidores y conserva '
  'el máximo. El cluster es el criterio que "suena más fuerte".')

h3('5.1 long_pump (impulso comprador) — suma lp')
table(['Componente', 'Condición', 'Puntos', 'Flag'], [
 ['Volumen', 'spike ≥ 10', '+45', 'extreme_volume_spike'],
 ['', 'spike ≥ 6', '+35', 'high_volume_spike'],
 ['', 'spike ≥ 3', '+25', 'volume_spike'],
 ['Precio 24h', '≥ 50%', '+35', 'price_parabolic'],
 ['', '≥ 25%', '+25', 'price_running'],
 ['', '≥ 10%', '+15', '—'],
 ['Imbalance', '≥ 0,80', '+20', 'bids_stacked'],
 ['', '≥ 0,65', '+10', '—'],
 ['Trampa', 'liq baja AND spike ≥ 3', '+15', 'low_liquidity_trap'],
], widths=[1.2, 1.9, 0.8, 2.0])

h3('5.2 classic (squeeze que va subiendo) — suma cl')
table(['Componente', 'Condición', 'Puntos'], [
 ['Imbalance', '≥0,80 → +40 · ≥0,70 → +30 · ≥0,60 → +18 · ≥0,55 → +8', ''],
 ['Precio 24h', '5–25% → +25 · 25–50% → +12', ''],
 ['Volumen', 'spike <3 → +15 · <6 → +8', ''],
 ['Trampa', 'liq baja AND imbalance ≥0,65 → +10', ''],
], widths=[1.3, 4.6, 0.3])

h3('5.3 Agregación')
code('score_long_pump = clamp(round(lp), 0, 100)\n'
     'score_classic   = clamp(round(cl), 0, 100)\n'
     'pump_score      = max(score_long_pump, score_classic)\n'
     'cluster = "classic" if score_classic > score_long_pump else "long_pump"')

h3('5.4 confidence_score (qué tan fiable es la señal)')
code('confidence = 35\n'
     'confidence += min(liquidity_usd / 10000, 35)   # +35 máx por libro profundo\n'
     'confidence += 15  si 10 <= cambio_24h <= 60     # move vivo, no ya dumpeado\n'
     'confidence += 10  si volume_spike >= 3\n'
     'confidence_score = clamp(round(confidence), 0, 100)   # rango ~35..95')
p('Idea: la confianza sube con liquidez REAL (difícil de falsear) y un move limpio aún no '
  'agotado. Los spikes de libro fino se quedan en confianza baja.', italic=True, color=GREY)

h3('5.5 classification (etiqueta)')
table(['Etiqueta', 'Condición (en orden)'], [
 ['criminal_pump_suspect', 'liq baja AND spike ≥ 6 AND cambio ≥ 10%'],
 ['active_pump', 'spike ≥ 6 AND cambio ≥ 25%'],
 ['accumulation_imbalance', 'imbalance ≥ 0,8 AND spike ≥ 3'],
 ['volume_anomaly', 'pump_score > 0'],
 ['no_signal', 'resto'],
 ['manipulation_suspect', 'sobrescribe si top-3 del libro > 80%'],
], widths=[2.2, 4.0])
p('liq baja = liquidity_usd < LOW_LIQUIDITY_USD (75.000).', size=9.5, color=GREY)

# ============================ 6 ============================
h1('6. Reglas de ENTRADA exactas')
p('Una entrada ocurre solo si TODOS estos gates pasan, en este orden (_auto_enter):')
table(['#', 'Gate', 'Condición para PASAR', 'Variable', 'Si falla'], [
 ['1', 'Confianza', 'confidence_score ≥ 50', 'ENTRY_MIN_CONFIDENCE', 'skip_low_confidence'],
 ['2', 'Anti-chase', 'price_change_pct_24h < 60', 'ENTRY_MAX_CHASE_PCT', 'skip_exhausted'],
 ['3', 'Volumen*', 'volume_spike ≥ 2,5', 'ENTRY_MIN_VOL_SPIKE', 'skip_low_volume'],
 ['4', 'Forensic', 'spread ≤ 1% · liq ≥ 120k · top3 ≤ 80%', 'FORENSIC_*', 'forensic_block'],
 ['5', 'Riesgo', 'caps RiskGuard (sección 13)', 'RiskLimits', 'rechazo'],
], widths=[0.4, 1.1, 2.3, 1.5, 1.4])
p('* El gate de volumen solo aplica al camino de escaneo. Las entradas del camino velocity '
  '(accel != None) ya demostraron aceleración real + precio subiendo, así que lo omiten.')
p('Además, para que un candidato llegue a evaluarse para entrada debe estar en '
  'waiting_confirmation: pump_score ≥ umbral adaptativo (70–90, arranca 75). Y el bot '
  'NO entra a un símbolo que ya tiene abierto.')

# ============================ 7 ============================
h1('7. Reglas de SALIDA exactas (PositionManager)')
p('Cada tick (15 s) llama pm.step(key, price, volume). Las salidas se evalúan en este '
  'orden, y la primera que dispara cierra (parcial o total):')
code(
'gain         = (price - entry) / entry * 100\n'
'drop_from_peak = (peak - price) / peak * 100\n'
'tick_drop    = (prev - price) / prev * 100\n'
'elapsed_min  = (now - entry_at) en minutos\n\n'
'1) HARD STOP   si gain <= -8%            -> vende 100% ("hard_stop")\n'
'2) DUMP        si tick_drop >= 10%       -> vende 100% ("dump")\n'
'3) BREAK-EVEN  arma cuando gain >= +4%; be_stop = entry*(1+0,5%)\n'
'               si armado y price <= be_stop -> vende 100% ("break_even")\n'
'4) TIME-STOP   si lateral (|gain| <= 3%) y volumen agotado -> vende 100% ("timeout")\n'
'5) TP1 (fase1) si gain >= +15%           -> vende 60%, pasa a fase 2 ("tp1")\n'
'6) TRAILING(f2) si drop_from_peak >= 10% -> vende el resto ("trailing")')

# ============================ 8 ============================
h1('8. Stop loss (hard stop)')
p('Es la PRIMERA comprobación (prioridad de protección de capital).', bold=True)
table(['Parámetro', 'Valor def.', 'Variable de entorno'], [
 ['HARD_STOP_PCT', '8%', 'PUMP_STOP_LOSS_PCT'],
], widths=[2.0, 1.2, 3.0])
code('if gain <= -HARD_STOP_PCT:  # gain <= -8%\n    vender 100% como "hard_stop"')
p('Nota: el LiveBroker además coloca un take-profit reduce-order al entrar (best-effort) y '
  'el stop_loss = entry*(1-8%) se adjunta a cada leg; en paper la protección efectiva es la '
  'lógica de pm.step. El Monte Carlo muestra que el hard_stop es el que absorbe los rugs '
  '(19,7% de los trades, -14,56% medio) — es el seguro contra la cola.')

# ============================ 9 ============================
h1('9. Trailing stop')
p('Activo en FASE 2 (después de tomar el TP1 parcial). Sigue el pico.')
table(['Parámetro', 'Valor def.', 'Variable'], [
 ['TRAIL_PCT', '10%', 'PUMP_TRAIL_PCT'],
], widths=[2.0, 1.2, 3.0])
code('peak_price se actualiza en cada tick si price > peak\n'
     'drop_from_peak = (peak_price - price) / peak_price * 100\n'
     'if fase == 2 and drop_from_peak >= 10%:  vender el resto como "trailing"')
p('En el Monte Carlo el trailing es el mejor reason: 7,5% de los trades, +14,02% medio, '
  '100% ganadores. Es donde se capturan los pumps reales que corren.', italic=True, color=GREEN)

# ============================ 10 ============================
h1('10. Take profits')
p('Toma parcial en dos fases (nunca ser la liquidez de salida).')
table(['Parámetro', 'Valor def.', 'Variable', 'Significado'], [
 ['TP1_PCT', '15%', 'PUMP_TP1_PCT', 'Gatillo de la toma parcial'],
 ['TP1_FRAC', '0,6 (60%)', 'PUMP_TP1_FRAC', 'Fracción vendida en fase 1'],
 ['TAKE_PROFIT_PCT', '25%', 'PUMP_TAKE_PROFIT_PCT', 'TP del leg en executor (live reduce-order)'],
], widths=[1.6, 1.1, 1.7, 1.8])
code('if fase == 1 and gain >= TP1_PCT:   # +15%\n'
     '    vender 60% como "tp1"; fase = 2  # el 40% restante corre con trailing')
p('TP1 se bajó de 30% a 15% porque el bot entra a mitad de move: un primer objetivo de '
  '+30% casi nunca disparaba y dejaba a los ganadores rascarse a break-even. Bancar 60% a '
  '+15% convierte pumps moderados en ganancias reales (Monte Carlo: +0,7%/trade de mejora).',
  italic=True, color=GREY)

# ============================ 11 ============================
h1('11. Dump detector y time-stop dinámico')
h3('Dump detector')
table(['Parámetro', 'Valor def.', 'Variable'], [
 ['DUMP_TICK_PCT', '10%', 'PUMP_DUMP_TICK_PCT'],
], widths=[2.0, 1.2, 3.0])
p('Una caída abrupta de un solo tick ≥ 10% = pánico de venta inmediato del resto. Atrapa '
  'el colapso súbito antes de que el hard_stop por gain acumulado lo vea.')

h3('Time-stop dinámico consciente del volumen (_time_stop_fires)')
p('Un move lateral NO se corta solo por ser lento: se corta cuando se queda sin combustible '
  '(volumen). "Vivo" = volumen 1m actual ≥ 50% del pico de volumen del trade.')
table(['Parámetro', 'Valor def.', 'Variable', 'Rol'], [
 ['TIMEOUT_BAND_PCT', '3%', 'PUMP_TIMEOUT_BAND_PCT', 'Lateral = |gain| ≤ 3%'],
 ['TIMEOUT_MINUTES', '8 min', 'PUMP_TIMEOUT_MINUTES', 'Mínimo para cortar un plano sin fuel'],
 ['VOLUME_ALIVE_FRAC', '0,5', 'PUMP_VOLUME_ALIVE_FRAC', 'Umbral "vivo" vs pico'],
 ['TIMEOUT_NO_VOL_MINUTES', '20 min', 'PUMP_TIMEOUT_NO_VOL_MINUTES', 'Fallback si no hay dato de volumen'],
 ['MAX_HOLD_MINUTES', '45 min', 'PUMP_MAX_HOLD_MINUTES', 'Tope duro aunque el volumen persista'],
], widths=[1.9, 1.0, 1.9, 1.6])
code('si |gain| > 3%:                          -> NO corta (deja correr TP/trail/stop)\n'
     'si hay volumen:\n'
     '    faded = last_vol < 0,5 * peak_vol\n'
     '    si faded y elapsed >= 8 min          -> corta ("timeout")\n'
     '    si elapsed >= 45 min                 -> corta (tope)\n'
     '    si vivo                              -> mantiene\n'
     'si NO hay volumen: corta si elapsed >= 20 min')
p('En el Monte Carlo el timeout es la salida más común (57,4%) con +0,93% medio y 71% de '
  'aciertos: saca capital de moves muertos sin perder, para reciclarlo.', italic=True, color=GREY)

# ============================ 12 ============================
h1('12. Gestión de capital')
table(['Parámetro', 'Valor def.', 'Variable', 'Significado'], [
 ['PAPER_BALANCE', '1000 USD', 'PUMP_PAPER_BALANCE', 'Balance demo por cuenta'],
 ['AUTO_ENTRY_USD', '100 USD', 'PUMP_AUTO_ENTRY_USD', 'Tamaño por entrada (por usuario)'],
 ['Allocation', 'mexc 100 / bitget 0', 'splits', 'Reparto del capital por venue'],
 ['per_leg', 'capital / nº exchanges', '—', 'Capital dividido por venue de ejecución'],
 ['SLIPPAGE_PCT', '0,5%', 'PUMP_PAPER_SLIPPAGE_PCT', 'Slippage simulado en paper'],
], widths=[1.5, 1.4, 1.6, 1.7])
p('Cada cuenta (UserBot) tiene su balance, posiciones, equity y P&L propios y aislados. El '
  'equity en paper = capital asignado + P&L realizado + P&L no realizado. Si la cuenta tiene '
  'claves de exchange, se muestra el balance REAL (read-only) en vez del paper.')
p('Multi-tenant: get_bot(uid) crea/recupera el bot de cada cuenta; los loops iteran sobre '
  'all_bots(). Mutar una cuenta nunca toca a otra (verificado en vivo).', italic=True, color=GREY)

# ============================ 13 ============================
h1('13. Gestión de riesgo (RiskGuard)')
p('Ningún orden — paper o live — se coloca sin pasar RiskGuard.evaluate(). Prefiere el Risk '
  'Engine canónico si RISK_ENGINE_URL está seteado, y para LIVE falla cerrado si ese '
  'servicio no responde. El espejo local aplica los mismos caps.')
table(['Límite', 'Valor def.', 'Variable', 'Bloqueo'], [
 ['max_position_size_usd', '500', '—', 'position size exceeds limit'],
 ['max_open_trades', '4', 'PUMP_MAX_OPEN_TRADES', 'open trade limit reached'],
 ['max_daily_loss_usd', '250', '—', 'daily loss limit reached'],
 ['max_drawdown_pct', '5%', '—', 'drawdown limit reached'],
 ['max_leverage', '2', '—', 'leverage exceeds limit'],
 ['kill_switch', 'off', '—', 'corta todo cuando se activa'],
], widths=[1.9, 1.0, 1.7, 1.8])
p('Estrategia "depredador": concentrada, no dispersa. Máximo 4 posiciones abiertas '
  'simultáneas (contadas de posiciones vivas, no de fills históricos). El kill switch es '
  'por usuario (cada cuenta puede cortar su propio bot).')

# ============================ 14 ============================
h1('14. Cooldowns')
table(['Cooldown', 'Valor def.', 'Variable', 'Efecto'], [
 ['Trigger velocity', '600 s', 'PUMP_VELOCITY_COOLDOWN_SECONDS', 'Tras disparar un símbolo, no vuelve a disparar 10 min'],
 ['Posición abierta', '—', '(pm.has)', 'No reentra a un símbolo que ya tiene abierto'],
 ['Umbral adaptativo', '—', '(_adaptive_threshold)', 'Sube tras pérdidas → menos entradas (cooldown implícito)'],
 ['Caché CoinGecko', '1800 s', 'TTL_SECONDS', 'No repregunta FDV/mcap del mismo símbolo en 30 min'],
 ['Token JWT grid', '12 h', '—', 'Cachea el SSO del grid por usuario'],
], widths=[1.7, 1.0, 2.1, 1.6])

# ============================ 15 ============================
h1('15. Filtros anti-rug (ForensicFilter)')
p('Honestidad primero: en un CEX se opera contra el libro del exchange, NO contra un '
  'contrato. Por eso los checks de DEX (honeypot, liquidez quemada, impuesto compra/venta, '
  'lista de holders) NO tienen fuente y NO se falsean. Lo medible del libro público sí se '
  'aplica en forensic_check():')
table(['Check', 'Umbral', 'Variable', 'Por qué'], [
 ['Spread máx', '≤ 1,0%', 'PUMP_FORENSIC_MAX_SPREAD_PCT', 'Libro fino sangra el spread al instante'],
 ['Liquidez mín', '≥ 120.000 USD', 'PUMP_FORENSIC_MIN_LIQUIDITY_USD', 'Los rugs son eventos de libro fino; libro profundo = mayor protección'],
 ['Concentración top-3', '≤ 80%', 'PUMP_FORENSIC_MAX_TOP_SHARE', 'Libro sostenido por pocas órdenes = proxy de actor único'],
], widths=[1.4, 1.3, 2.2, 1.5])
p('El piso de liquidez se subió de 50k a 120k porque el Monte Carlo demostró que TODA la '
  'expectativa negativa viene de la cola de rugs/gap-down, y los rugs son eventos de libro '
  'fino. Subir el piso = muchos menos trades pero EV materialmente mejor.', italic=True, color=GREY)
p('Capas anti-rug adicionales (no en forensic): dump detector (-10% en un tick), hard_stop '
  '(-8%), y top_book_share > 80% marca manipulation_suspect en el escaneo.')

# ============================ 16 ============================
h1('16. Sistema de aprendizaje')
p('Hay DOS mecanismos de aprendizaje, separados:')

h3('16.1 Umbral adaptativo loss-averse (_apply_learning)')
code('al cerrar un trade con calidad y pnl:\n'
     '  si pnl < 0:                 umbral = min(90, umbral + 2)   # más selectivo tras pérdida\n'
     '  elif calidad == "early_entry": umbral = max(70, umbral - 1)  # pump real cazado temprano\n'
     '  (una ganancia pequeña no-early deja el umbral igual)')
p('Banda 70–90, arranca en 75. REEMPLAZA al loop viejo que bajaba el umbral tras trades '
  '"tardíos" — aquel entraba en espiral: cada timeout de pérdida pequeña lo hacía menos '
  'selectivo, compraba más ruido, más timeouts. Ahora una pérdida lo hace MÁS exigente.',
  italic=True, color=GREY)
h3('Calidad de entrada (entry_quality)')
table(['Calidad', 'Condición'], [
 ['late_entry', 'peak_gain < 5% O segundos-al-pico < 60 (compró cerca del techo)'],
 ['early_entry', 'peak_gain ≥ 30%'],
 ['perfect_entry', 'resto'],
], widths=[1.6, 4.6])

h3('16.2 LearningLab (analítica del cerebro)')
p('Registra cada token alertado y sigue su MFE/MAE, lead-time (¿la alerta fue ANTES del '
  'pump?), precisión/recall y la contribución de cada componente del score. Propone ajustes '
  'de umbral una vez los resultados se asientan. El usuario puede reportar pumps no '
  'detectados (/learning/missed) para bajar el recall. Endpoint /learning expone el snapshot.')

# ============================ 17 ============================
h1('17. Todas las fórmulas')
h3('Escaneo / score')
code('volume_spike    = vol(última 1h cerrada) / media(vol 1h previas)   # 1.0 = sin spike\n'
     'imbalance       = bid_notional / (bid_notional + ask_notional)     # dentro de ±2% del mid\n'
     'liquidity_usd   = bid_notional + ask_notional   (dentro de ±2% del mid)\n'
     'spread_pct      = (best_ask - best_bid) / best_ask * 100\n'
     'top_book_share  = sum(top-3 bid notional) / sum(bid notional)\n'
     'confidence      = 35 + min(liq/10000, 35) + (15 si 10<=chg<=60) + (10 si spike>=3)')
h3('Velocity')
code('accel = vol_1m(última cerrada) / media(vol_1m previas)\n'
     'dispara si  accel >= 4  AND  last_close >= prev_close  AND  vela nueva  AND  cooled(600s)\n'
     'baseline_vol = baseline_vol*0,8 + base*0,2   # EMA lenta del baseline')
h3('Ejecución (paper)')
code('fill_price = entry * (1 + 0,5%/100)  para compra   (slippage)\n'
     'amount     = notional_usd / fill_price\n'
     'sl         = entry * (1 - 8%/100)\n'
     'tp         = entry * (1 + 25%/100)')
h3('Salida / P&L')
code('gain           = (price - entry)/entry * 100\n'
     'drop_from_peak = (peak - price)/peak * 100\n'
     'tick_drop      = (prev - price)/prev * 100\n'
     'be_stop        = entry * (1 + 0,5%/100)\n'
     'pnl(parcial)   = (price - entry) * sell_qty\n'
     'pnl_pct(cierre) = realized_pnl / (entry * initial_qty) * 100')
h3('Equity / P&L de cuenta')
code('paper_equity = capital_asignado + sum(realized) + sum(unrealized abiertas)\n'
     'pnl_7d       = realized(últimos 7d) + unrealized abiertas')

# ============================ 18 ============================
h1('18. Todos los parámetros configurables (variables de entorno)')
h3('Escaneo y cerebro')
table(['Variable', 'Def.', 'Descripción'], [
 ['PUMP_EXEC_MODE', 'paper', 'paper | live'],
 ['PUMP_SCAN_EXCHANGES', 'binance,mexc,bitget', 'Exchanges a escanear'],
 ['PUMP_SCAN_INTERVAL_SECONDS', '300', 'Cadencia del escaneo'],
 ['PUMP_VELOCITY_TICK_SECONDS', '10', 'Cadencia del watcher rápido'],
 ['PUMP_VELOCITY_ACCEL_FACTOR', '4', 'Aceleración de volumen que dispara'],
 ['PUMP_VELOCITY_WATCH_TOP_N', '8', 'Símbolos en la hot-list'],
 ['PUMP_VELOCITY_WATCH_MIN_SCORE', '40', 'Score mínimo para vigilar'],
 ['PUMP_VELOCITY_COOLDOWN_SECONDS', '600', 'Cooldown por símbolo'],
 ['PUMP_THRESHOLD_FLOOR / _CEIL', '70 / 90', 'Banda del umbral adaptativo'],
], widths=[2.6, 1.1, 2.5])
h3('Entrada')
table(['Variable', 'Def.', 'Descripción'], [
 ['PUMP_AUTO_ENTRY', 'true', 'Auto-entrada (solo paper)'],
 ['PUMP_AUTO_ENTRY_USD', '100', 'USD por entrada'],
 ['PUMP_ENTRY_MIN_CONFIDENCE', '50', 'Piso de confianza'],
 ['PUMP_ENTRY_MAX_CHASE_PCT', '60', 'Máximo % 24h para no perseguir el techo'],
 ['PUMP_ENTRY_MIN_VOL_SPIKE', '2.5', 'Spike mínimo (camino scan)'],
], widths=[2.6, 1.1, 2.5])
h3('Forensic / anti-rug')
table(['Variable', 'Def.', 'Descripción'], [
 ['PUMP_FORENSIC_MAX_SPREAD_PCT', '1.0', 'Spread máximo'],
 ['PUMP_FORENSIC_MIN_LIQUIDITY_USD', '120000', 'Liquidez mínima'],
 ['PUMP_FORENSIC_MAX_TOP_SHARE', '0.80', 'Concentración top-3 máxima'],
], widths=[2.6, 1.1, 2.5])
h3('Salida (PositionManager)')
table(['Variable', 'Def.', 'Descripción'], [
 ['PUMP_TP1_PCT', '15', 'Gatillo toma parcial'],
 ['PUMP_TP1_FRAC', '0.6', 'Fracción vendida en TP1'],
 ['PUMP_TRAIL_PCT', '10', 'Trailing desde el pico'],
 ['PUMP_STOP_LOSS_PCT', '8', 'Hard stop'],
 ['PUMP_DUMP_TICK_PCT', '10', 'Dump de un tick'],
 ['PUMP_TIMEOUT_MINUTES', '8', 'Mín. para cortar plano sin fuel'],
 ['PUMP_TIMEOUT_BAND_PCT', '3', 'Banda lateral'],
 ['PUMP_BREAKEVEN_PCT', '4', 'Gain que arma break-even'],
 ['PUMP_BREAKEVEN_MARGIN_PCT', '0.5', 'SL por encima de entry'],
 ['PUMP_VOLUME_ALIVE_FRAC', '0.5', 'Umbral volumen vivo vs pico'],
 ['PUMP_TIMEOUT_NO_VOL_MINUTES', '20', 'Fallback sin volumen'],
 ['PUMP_MAX_HOLD_MINUTES', '45', 'Tope duro de hold'],
 ['PUMP_TAKE_PROFIT_PCT', '25', 'TP del leg (live reduce-order)'],
], widths=[2.6, 1.1, 2.5])
h3('Capital y riesgo')
table(['Variable', 'Def.', 'Descripción'], [
 ['PUMP_PAPER_BALANCE', '1000', 'Balance demo por cuenta'],
 ['PUMP_PAPER_SLIPPAGE_PCT', '0.5', 'Slippage paper'],
 ['PUMP_MAX_OPEN_TRADES', '4', 'Posiciones abiertas máx'],
 ['PUMP_EXEC_EXCHANGES', 'mexc,bitget', 'Venues de ejecución'],
 ['RISK_ENGINE_URL', '(none)', 'Servicio de riesgo canónico opcional'],
], widths=[2.6, 1.1, 2.5])
h3('Claves (solo live / balance real — sin permiso de retiro)')
table(['Variable', 'Uso'], [
 ['MEXC_API_KEY / MEXC_SECRET', 'Spot MEXC'],
 ['BITGET_API_KEY / BITGET_SECRET / BITGET_PASSWORD', 'Spot Bitget'],
 ['BINANCE_API_KEY / BINANCE_SECRET', 'Spot Binance'],
 ['SUPABASE_* , APP_PASSWORD, TELEGRAM_*', 'Persistencia / auth / notificaciones (.env)'],
], widths=[3.6, 2.6])

# ============================ 19 ============================
h1('19. Ejemplo real: una señal de principio a fin + Monte Carlo')
p('Ejemplo ILUSTRATIVO con los parámetros reales del bot (precios de ejemplo para mostrar '
  'la mecánica; no son datos de mercado reales).', italic=True, color=GREY)
h3('19.1 Detección')
table(['Campo', 'Valor de ejemplo'], [
 ['Símbolo', 'ABC/USDT en mexc'],
 ['price_change_pct_24h', '+38%'],
 ['volume_spike', '7,0x'],
 ['imbalance', '0,72'],
 ['liquidity_usd', '180.000'],
 ['spread_pct', '0,4%'],
 ['top_book_share', '0,55'],
], widths=[2.0, 3.0])
h3('19.2 Score')
code('long_pump: vol 7x(+35) + chg 38%(+25=price_running) + imbalance 0,72(+10) = 70\n'
     'classic:   imbalance 0,72(+30) + chg 25-50(+12) + vol<? no = 42\n'
     'pump_score = max(70, 42) = 70   cluster = long_pump\n'
     'confidence = 35 + min(180000/10000,35)=35 + 15(chg in 10..60) + 10(spike>=3) = 95\n'
     'classification = active_pump  (spike>=6 AND chg>=25)')
p('pump_score 70 < umbral 75 → de momento "watching". Si el move sigue y el velocity ve '
  'accel ≥ 4x, dispara igual por el camino rápido. Supongamos que en el siguiente escaneo '
  'el score sube a 78 ≥ 75 → waiting_confirmation → se evalúa la entrada.')
h3('19.3 Gates de entrada')
code('(1) confidence 95 >= 50            OK\n'
     '(2) chg 38% < 60%                  OK\n'
     '(3) volume_spike 7,0 >= 2,5        OK\n'
     '(4) forensic: spread 0,4<=1, liq 180k>=120k, top3 0,55<=0,80   OK\n'
     '(5) RiskGuard: <4 abiertas, size 100<500   OK  -> ENTRADA')
p('Entrada paper: fill = precio * 1,005 (slippage). Compra de 100 USD. pm.open() registra la '
  'posición y arranca el motor de salida.')
h3('19.4 Gestión hasta el cierre (un desenlace posible)')
table(['t', 'Precio vs entry', 'Evento'], [
 ['+2 min', '+5%', 'arma break-even (be_stop = entry+0,5%)'],
 ['+6 min', '+16%', 'TP1: vende 60% a +15% efectivo; pasa a fase 2'],
 ['+9 min', 'pico +22%', 'peak_price actualizado'],
 ['+12 min', '+11% (cae 11% desde pico)', 'TRAILING: vende el 40% restante'],
 ['cierre', '—', 'pnl_pct ≈ +15,4%; calidad = perfect_entry'],
], widths=[1.0, 2.2, 2.6])
p('Como cerró en ganancia y no fue early_entry estricto, el umbral queda igual. Si hubiera '
  'cerrado en pérdida (p. ej. rug → hard_stop a -8%), el umbral subiría a 77 (más exigente).')

h3('19.5 Resultado agregado — Monte Carlo 1.000.000 de casos')
p('Se corrió la lógica REAL del bot (mismos gates + mismo PositionManager) sobre un millón '
  'de escenarios sintéticos calibrados, con coste 0,20%/trade ida y vuelta.')
table(['Métrica', 'Con gates (lo que opera el bot)', 'Sin gates (toda señal)'], [
 ['Entradas', '20.788 (2,1% de señales)', '1.000.000'],
 ['Expectativa/trade', '-1,23%', '-5,50%'],
 ['Win rate', '52,7%', '48,3%'],
 ['Profit factor', '0,63', '0,25'],
 ['Ganancia media', '+3,98%', '—'],
 ['Pérdida media', '-7,03%', '—'],
], widths=[1.8, 2.4, 2.0])
p('Los gates aportan +4,27%/trade de expectativa. Aun así el total es negativo.', bold=True, color=RED)
h3('Desglose por motivo de salida (con gates)')
table(['Motivo', '% trades', 'Media', 'Win %'], [
 ['timeout', '57,4%', '+0,93%', '71,2%'],
 ['hard_stop', '19,7%', '-14,56%', '0,0%'],
 ['break_even', '14,3%', '-1,04%', '21,7%'],
 ['trailing', '7,5%', '+14,02%', '100%'],
 ['tp1', '1,1%', '+17,43%', '100%'],
 ['open_eod', '0,1%', '+9,09%', '100%'],
], widths=[1.6, 1.2, 1.4, 1.4])
p('Diagnóstico: TODA la pérdida es la cola de rugs (hard_stop, -14,56% medio). Los pumps '
  '"criminales" de libro fino que el bot detecta son justamente los más propensos a rug. El '
  'filtrado es el verdadero valor del bot; el auto-trading no voltea a positivo. '
  'Recomendación: mantener PAPER; usar el bot como radar de detección/evitación.',
  bold=True)

# ============================ 20 ============================
h1('20. Por qué existe cada regla')
table(['Regla', 'Por qué existe'], [
 ['Excluir majors/stables/leveraged', 'No son objetivos de pump criminal; ensucian el universo.'],
 ['Volumen 24h en [100k, 60M]', 'Bajo = no operable; alto = blue chip que no se "pumpea".'],
 ['Solo gainers (chg>0)', 'Un pump por definición sube; mirar caídas es otro problema.'],
 ['Dos criterios (long_pump/classic)', 'Distinguir impulso comprador de squeeze que va subiendo.'],
 ['confidence_score', 'Separar señales con liquidez real (fiables) de spikes de libro fino.'],
 ['Gate confianza ≥ 50', 'Las señales de libro fino solo sangran el spread; se descartan.'],
 ['Gate anti-chase < 60%', 'Comprar algo ya +60% = comprar el techo del blow-off.'],
 ['Gate volumen ≥ 2,5x', 'Sin volumen real detrás, el "move" es ruido.'],
 ['ForensicFilter (spread/liq/top3)', 'Lo único auditable de manipulación en un CEX; sin inventar on-chain.'],
 ['Liquidez ≥ 120k', 'Los rugs son eventos de libro fino; libro profundo = la mayor protección.'],
 ['Hard stop -8%', 'Tope de pérdida por trade; absorbe la cola de rugs.'],
 ['Dump detector -10% tick', 'Atrapa el colapso súbito antes que el stop por gain acumulado.'],
 ['Break-even a +4%', 'Una vez en verde, no devolver la ganancia: bloquea a ~breakeven.'],
 ['Time-stop por volumen', 'No castigar un lateral con fuel; sí liberar capital de moves muertos.'],
 ['TP1 60% a +15%', 'Asegurar capital pronto (el bot entra a mitad de move); no ser exit-liquidity.'],
 ['Trailing 10% en fase 2', 'Dejar correr al 40% restante en los pumps reales, con give-back acotado.'],
 ['Máx 4 abiertas', 'Estrategia concentrada (depredador), no dispersa.'],
 ['Cooldown velocity 600s', 'Evitar disparos repetidos del mismo símbolo en segundos.'],
 ['No reentrar símbolo abierto', 'Evita duplicar exposición al mismo riesgo.'],
 ['Umbral loss-averse', 'Romper la espiral: tras una pérdida, ser MÁS selectivo, no menos.'],
 ['Paper por defecto', 'Sin edge demostrado no se arriesga dinero real; live es opt-in con claves propias.'],
 ['RiskGuard falla cerrado (live)', 'Si el servicio de riesgo no responde, mejor no operar que operar a ciegas.'],
], widths=[2.2, 4.0])

# Cierre
doc.add_paragraph()
p('— Fin del documento. Generado del código fuente real de apps/pump-reader. '
  'Modo por defecto: PAPER. Nada de datos de negocio inventados.', italic=True, color=GREY, size=9)

out = r'C:\Users\osval\OneDrive\Escritorio\TradeOS_Amatista_Documentacion_Tecnica.docx'
doc.save(out)
print('SAVED', out)
print('paragraphs', len(doc.paragraphs), 'tables', len(doc.tables))
