# -*- coding: utf-8 -*-
"""Informe: Auditoria estructural + rediseno de TradeOS Pump Reader.
Contenido = analisis de codigo real (no inventado). Honestidad sobre datos."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INK = RGBColor(0x0c, 0x10, 0x18)
PINK = RGBColor(0xff, 0x2f, 0x6e)
GREY = RGBColor(0x55, 0x5b, 0x66)
GREEN = RGBColor(0x18, 0x8a, 0x5a)
RED = RGBColor(0xc0, 0x2a, 0x2a)
AMBER = RGBColor(0xb5, 0x6a, 0x00)

doc = Document()
base = doc.styles['Normal']; base.font.name = 'Calibri'; base.font.size = Pt(10.5)

def _shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)

def h1(t):
    p = doc.add_heading(t, level=1)
    for r in p.runs: r.font.color.rgb = PINK
    return p

def h2(t):
    p = doc.add_heading(t, level=2)
    for r in p.runs: r.font.color.rgb = INK
    return p

def h3(t):
    p = doc.add_heading(t, level=3)
    for r in p.runs: r.font.color.rgb = GREY
    return p

def p(t, bold=False, italic=False, color=None, size=10.5):
    par = doc.add_paragraph(); r = par.add_run(t)
    r.bold = bold; r.italic = italic; r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    return par

def callout(t, color=RED):
    par = doc.add_paragraph(); par.paragraph_format.left_indent = Inches(0.15)
    par.paragraph_format.space_before = Pt(4); par.paragraph_format.space_after = Pt(6)
    r = par.add_run(t); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = color
    return par

def bullet(t, prefix=None):
    par = doc.add_paragraph(style='List Bullet')
    if prefix:
        r = par.add_run(prefix); r.bold = True; par.add_run(t)
    else: par.add_run(t)
    return par

def code(t):
    par = doc.add_paragraph(); par.paragraph_format.left_indent = Inches(0.2)
    par.paragraph_format.space_before = Pt(2); par.paragraph_format.space_after = Pt(6)
    for line in t.split('\n'):
        r = par.add_run(line + '\n'); r.font.name = 'Consolas'; r.font.size = Pt(8.6)
        r.font.color.rgb = RGBColor(0x1b, 0x2a, 0x4a)
    return par

def table(headers, rows, widths=None):
    tb = doc.add_table(rows=1, cols=len(headers)); tb.style = 'Light Grid Accent 1'
    tb.alignment = WD_TABLE_ALIGNMENT.LEFT
    hc = tb.rows[0].cells
    for i, htext in enumerate(headers):
        hc[i].text = ''; rr = hc[i].paragraphs[0].add_run(htext)
        rr.bold = True; rr.font.size = Pt(9); rr.font.color.rgb = RGBColor(0xff,0xff,0xff)
        _shade(hc[i], '2f3a52')
    for row in rows:
        cells = tb.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''; rr = cells[i].paragraphs[0].add_run(str(val)); rr.font.size = Pt(8.6)
    if widths:
        for i, w in enumerate(widths):
            for row in tb.rows: row.cells[i].width = Inches(w)
    doc.add_paragraph(); return tb

# ===== PORTADA =====
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('Amatista · TradeOS'); r.bold = True; r.font.size = Pt(28); r.font.color.rgb = PINK
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run('Auditoría Estructural y Rediseño de Arquitectura'); r.font.size = Pt(15); r.font.color.rgb = INK
s2 = doc.add_paragraph(); s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s2.add_run('¿El bot persigue el fenómeno correcto? — Momentum Detector → Early Preparation Detector')
r.italic = True; r.font.size = Pt(10.5); r.font.color.rgb = GREY
m = doc.add_paragraph(); m.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = m.add_run('Rol: Principal Quant Engineer · Market microstructure · 2026-06-18 · Modo: PAPER')
r.font.size = Pt(9); r.font.color.rgb = GREY

p('')
callout('VEREDICTO EN UNA LÍNEA: el sistema es un detector de MOMENTUM VISIBLE EN CURSO (entra DESPUÉS '
        'de que el pump empezó). No detecta preparación. El fenómeno que persigue es estructuralmente '
        'negativo porque el trigger del bot y el trigger del rug son el mismo evento. La rama de '
        '"acumulación" nunca se probó: faltan los datos. Recomendación: construir solo el recolector '
        'de datos y dejar que los datos decidan. Mantener PAPER.', RED)

doc.add_page_break()

# ===== ADVERTENCIA EPISTEMICA =====
h1('0. Advertencia epistémica (leer primero)')
p('Un hecho condiciona TODO el informe y ningún rediseño puede saltárselo:')
callout('El sistema nunca ha grabado la microestructura pre-pump de tokens reales.', AMBER)
bullet('Cada escaneo hace _candidates.clear() (main.py:994) y reescribe un snapshot. No hay serie temporal de imbalance / liquidez / profundidad.', None)
bullet('El único EV = -1,23% viene de tools/simulate.py: un generador SINTÉTICO cuyas correlaciones fueron definidas a mano (rug ∝ 1/liquidez; features ∝ calidad latente + ruido).', None)
p('Consecuencia dura para las tareas empíricas (ablation, edge-search, estadística de hard-stops, '
  'universo A/B/C):', bold=True)
callout('Correr esos estudios sobre el simulador MIDE LOS SUPUESTOS DEL SIMULADOR, NO EL MERCADO. '
        'Sería circular. No se lavan números sintéticos como descubrimientos.', RED)
p('Qué SÍ es evidencia real: el código fuente. Las Tareas 1, 2 y el diagnóstico se responden con '
  'código verificable (archivo:línea). Las tareas empíricas NO son contestables hoy con rigor — '
  'requieren el pipeline de datos (Etapa 1). Se marca explícitamente en cada punto.', italic=True, color=GREY)

# ===== TAREA 1 =====
h1('1. Auditoría: ¿qué detecta realmente?')
p('Respuesta con evidencia de código, no opiniones.')
table(['#', 'Pregunta', 'Veredicto', 'Evidencia (archivo:línea)'], [
 ['1', 'Acumulación temprana', 'NO', 'Universo filtra "if change<=0: continue" y ordena por % desc, top-20 (scanner.py:375-381). Token lateral (chg≈0) nunca entra al shortlist.'],
 ['2', 'Preparación de pump', 'NO', 'Score premia precio YA movido: +15/+25/+35 por chg ≥10/≥25/≥50% (scanner.py:215). Confidence exige 10≤chg≤60 (scanner.py:265).'],
 ['3', 'Persistencia temporal', 'NO', 'score_candidate recibe escalares de UN snapshot. Sin ventana, sin buffer, sin Δ. clear() cada pasada (main.py:994). Evalúa fotos, no películas.'],
 ['4', 'Intención institucional/coordinada', 'NO', 'Único proxy: top_book_share (top-3 del bid) en un snapshot (scanner.py:136-155). No sigue reposición de órdenes, iceberg ni absorción en el tiempo.'],
 ['5', 'Manipulación', 'PARCIAL (estática)', 'top_book_share>0.80 → flag manipulation_suspect (scanner.py:315). Umbral instantáneo, no patrón (spoofing requiere tiempo, no medido).'],
 ['6', 'Rugs', 'NO (solo reacciona)', 'Cero predicción. Se "detecta" cuando ya pasó: hard_stop -8% (position_manager.py:129), dump -10%/tick (:133). Pre-trade solo piso de liquidez estático.'],
 ['7', 'Precio simplemente moviéndose', 'SÍ — esto es lo que detecta', 'Universo=gainers, score∝chg, velocity exige rising (velocity.py:133), gate exige momentum+spike. Tres capas convergen en "precio ya sube con volumen".'],
], widths=[0.3, 1.7, 1.3, 3.0])
callout('CONCLUSIÓN: detector de momentum visible en curso. 6 de 7 fenómenos sofisticados: ausentes o reactivos.')

# ===== TAREA 2 =====
h1('2. Errores de diseño: ¿entra antes o después?')
h3('2.1 Dependencia del precio (auditada)')
p('Cuánto del pump_score proviene del precio (scanner.py:188-282):')
bullet('Rama long_pump (dominante): el precio aporta hasta 35/100 directo.', None)
bullet('Pero el efecto real es mayor: el precio es GATE de universo (sin chg>0 no existes) y GATE de confidence (sin 10≤chg≤60, confidence ~35-50, bajo el piso 50).', None)
callout('El precio no "pesa 35%": es CONDICIÓN NECESARIA DE EXISTENCIA. Peso efectivo ≈ 100%. '
        'Contamina la señal en los tres niveles: universo, score y trigger.', RED)
p('Velocity añade segunda capa de precio: rising = last_close >= prev_close es obligatorio para '
  'disparar (velocity.py:133-136). El anti-chase (chg<60) no reduce dependencia — solo recorta la '
  'cola alta; sigue exigiendo que el move empezó.')

h3('2.2 ¿A (pump ya empezó) o B (antes)?')
callout('Respuesta: A, sin ambigüedad.', RED)
bullet('Universo: solo tokens YA verdes, ordenados por cuánto ya subieron (scanner.py:376-381).', '1. ')
bullet('Confirmación: pump_score ≥ 75 requiere chg + spike YA presentes.', '2. ')
bullet('Velocity: dispara en la IGNICIÓN (accel≥4x mientras sube), no en la carga.', '3. ')
bullet('El propio gate anti-chase chg<60% (main.py:898) es una confesión del código: "intentamos no entrar demasiado tarde" → admite que la entrada es tardía por diseño.', '4. ')
p('El bot entra en la ventana [inicio visible, +60%] del move. Es momentum-chasing: la peor zona '
  'riesgo/recompensa en microcaps. Ya pagaste el spread del libro fino y compraste a quien acumuló barato.')

# ===== HARD STOPS =====
h1('3. Análisis especial: Hard Stops')
p('Lo honesto: no se puede contestar con datos hoy.', bold=True)
bullet('En el sim, los hard_stop SON el arquetipo "rug", y rug está cableado a baja liquidez (simulate.py:71-119). El sim dirá "hard stops = baja liquidez" PORQUE ASÍ SE ESCRIBIÓ. Circular.', None)
bullet('En la realidad: DESCONOCIDO. Ni un solo hard_stop real grabado con su contexto de libro. insert_exit guarda symbol/price/pnl/reason (store.py:118), NO la microestructura de la entrada ni la serie previa.', None)
callout('Lo que el DISEÑO garantiza sin datos: como entras en momentum sobre libros finos, tu '
        'distribución de salidas está sesgada a la cola izquierda. El rug y tú compartís el MISMO '
        'trigger (volumen+precio explotando en poca liquidez). Por construcción, los hard stops y '
        'tus "mejores señales" son el mismo objeto. Patología arquitectónica, no de parámetros.', RED)
p('Para responder de verdad "¿qué distingue a los hard stops?": grabar entry-context + serie '
  'pre-entrada de cada trade, acumular N≥200 reales, correr el discriminante. Es Etapa 1+4 del plan.',
  italic=True, color=GREY)

# ===== REDISENO =====
h1('4. Rediseño: Momentum Detector → Early Preparation Detector')
callout('Principio rector: el fenómeno objetivo es la divergencia VOLUMEN/LIQUIDEZ SIN precio. '
        'Acumulación = el libro se profundiza y el volumen sube MIENTRAS el precio se queda quieto. '
        'Es lo contrario de lo que el bot busca hoy.', INK)
p('Los 10 cambios pedidos colapsan en 4 piezas estructurales + 1 precondición de datos. No son 10 '
  'parches: son una arquitectura.')

h3('Precondición (Etapa 1, NO opcional): TimeSeriesStore + Recorder')
p('Sin esto, los Cambios #2/#3/#6 no tienen materia prima y los #9/#10 son imposibles. Hoy no existe. '
  'Las funciones de medición YA existen (_orderbook_metrics, _forensic_metrics, fetch_1m_volume); solo '
  'no se guarda la secuencia.')
code('# app/microstructure.py (NUEVO) — buffer rolling por simbolo\n'
     '@dataclass\n'
     'class MicroSnapshot:\n'
     '    t; price; vol_1m; imbalance; liquidity_usd; spread_pct; top_book_share\n'
     '    bid_depth_usd; ask_depth_usd   # NUEVO\n'
     '@dataclass\n'
     'class SymbolSeries:\n'
     '    buf: deque(maxlen=120)         # 120 intervalos\n'
     '# loop _observe_loop (~60s): toma MicroSnapshot de cada simbolo en watchlist y lo persiste')

h3('Cambio #1 — Score con precio mínimo')
p('El precio pasa de "requisito" a "descuento por tardanza", invirtiendo su rol:')
code('prep = 0.55*accumulation + 0.30*persistence + 0.15*sequence\n'
     'prep *= (1 - clamp(price_change_1h/40, 0, 0.6))   # cuanto mas ya subio, menos vale')

h3('Cambio #2 — AccumulationScore (0-100): detecta acumulación silenciosa')
code('def accumulation_score(w):\n'
     '    s = 0\n'
     '    if slope(vol_1m)>0 and abs(pct_change(price))<2: s += 30   # vol sube, precio plano = la firma\n'
     '    if slope(bid_depth_usd)>0:                        s += 20   # alguien apila bids\n'
     '    s += 20 * absorption_ratio(w)                                # absorbe ventas (0..1)\n'
     '    if slope(spread_pct)<0:                           s += 15   # spread comprimiendose\n'
     '    if 0.60 <= mean(imbalance) <= 0.85:               s += 15   # comprador sin blow-off\n'
     '    return min(100, s)')

h3('Cambio #3 + #6 — PersistenceScore y SequenceEngine (unificados)')
p('Persistencia = secuencia sostenida. "vol 1.5→2→3→5x vale más que 5x ahora" = monotonía + aceleración.')
code('def persistence_score(w):\n'
     '    s  = 25*frac_intervals(w, imbalance>0.65)   # imbalance sostenido\n'
     '    s += 25*monotonic_up(vol_1m)                 # vol creciente monotono\n'
     '    s += 25*monotonic_up(liquidity_usd)          # liq creciente\n'
     '    s += 25*sustained_absorption(w)              # absorcion sostenida\n'
     '    return min(100, s)\n'
     'def sequence_bonus(vols):                         # premia la FORMA (rampa) sobre el nivel\n'
     '    ratios = [b/a for a,b in zip(vols, vols[1:])]\n'
     '    return clamp(mean(ratios)-1,0,1) * (1 if is_increasing(ratios) else 0.3)')

h3('Cambio #4 — Fase de observación real (state machine)')
p('Prohibir entrada inmediata. Hoy _perform_scan entra en la misma pasada que detecta (main.py:1029-1034).')
code('SCAN -> CANDIDATE -> WATCHLIST -> [OBSERVE 15/30/60m grabando serie] -> CONFIRMATION -> ENTRY')
p('Confirmación = accumulation y persistence se MANTUVIERON altos durante la ventana Y aún no hubo '
  'blow-off. Qué ventana (15/30/60) da más edge: NO se puede decidir hoy — es un parámetro a barrer '
  'DESPUÉS de tener eventos reales etiquetados, no antes.')

h3('Cambio #5 — Separar PumpScore y RugRiskScore (dos modelos, nunca uno)')
p('forensic_check (gate binario) → reemplazado por rug_risk_score continuo con deterioro temporal:')
code('def rug_risk_score(w):                                # probabilidad de COLAPSO, independiente del pump\n'
     '    r = 0\n'
     '    if slope(liquidity_usd)<0:  r += 25              # liquidez desapareciendo\n'
     '    if slope(spread_pct)>0:     r += 20              # spread ensanchando\n'
     '    if mean(top_book_share)>0.8:r += 20              # concentracion\n'
     '    r += 25*violent_moves_without_support(w)         # velas sin profundidad detras\n'
     '    if bid_pulled_suddenly(w):  r += 10              # retirada de bid (spoof/exit)\n'
     '    return min(100, r)\n'
     '# decision: entrar si pump>=P AND rug<=R  (matriz, planos separados, NO suma)')

h3('Cambios #7/#8/#9/#10 — son experimentos, no código (y hoy no corren)')
table(['Cambio', 'Estado', 'Detalle'], [
 ['#7 Universo A/B/C', 'Requiere Etapa 1', 'B (vol+lateral) y C (imbalance persistente) no existen como universo hasta grabar series. A (gainers) ya es −EV direccional; B/C desconocidos.'],
 ['#8 Lead time', 'Infra parcial YA existe', 'LearningLab.lead_secs() (learning.py:67) mide alert→peak, no detección→inicio-de-move, y no alimenta el gate. Reutilizable, hay que reorientarla.'],
 ['#9 Ablation', 'NO ejecutable hoy', 'Sobre el sim = circular. Sobre datos reales = imposible (no hay datos). Solo tras Etapa 1+4.'],
 ['#10 Edge search', 'NO ejecutable hoy', 'Idem. Buscar regiones +EV sobre datos sintéticos es buscar lo que ya se metió en el generador.'],
], widths=[1.4, 1.4, 3.4])

# ===== ENTREGA FINAL =====
h1('5. Entrega final')
h3('5.1 Diagnóstico / módulos mal enfocados')
table(['Módulo', 'Estado', 'Diagnóstico'], [
 ['scanner._is_altcoin + universo', 'mal enfocado', 'Filtra por gainer → ciego a la preparación'],
 ['scanner.score_candidate', 'mal enfocado', 'Score ∝ precio; sin tiempo; conflación pump+rug'],
 ['velocity', 'mal enfocado', 'rising obligatorio → detecta ignición, no carga'],
 ['_perform_scan (entrada inmediata)', 'error estructural', 'No hay fase de observación'],
 ['forensic_check', 'insuficiente', 'Gate binario estático, no RugRiskScore temporal'],
 ['position_manager', 'CORRECTO', 'Motor de salida sólido (TP1/trailing/break-even/time-stop)'],
], widths=[2.1, 1.2, 2.9])

h3('5.2 Módulos que SÍ aportan edge (no eliminar)')
bullet('position_manager: el filtrado/salidas dan +4,27%/trade (medido en sim). Bien diseñado.', None)
bullet('Funciones de medición (_orderbook_metrics, _forensic_metrics, fetch_1m_volume): correctas y REUTILIZABLES para el recolector.', None)
bullet('LearningLab: analítica de lead-time/MFE/MAE/precision/recall, reorientable.', None)

h3('5.3 Eliminar / Crear')
p('Eliminar: nada se borra, se REUBICA. Invertir/ampliar "change<=0 continue" (incluir laterales). '
  'Quitar el aporte de precio a confidence. Partir pump_score único en pump/rug.', None)
p('Crear: microstructure.py (serie), accumulation.py, persistence.py, sequence.py, rug_risk.py, '
  'state-machine de observación en main.py, loop _observe_loop, tabla micro_series.', None)

h3('5.4 Archivos y código exactos a modificar')
table(['Archivo:línea', 'Cambio'], [
 ['scanner.py:376', 'Quitar "if change<=0: continue"; ampliar universo a laterales con volumen'],
 ['scanner.py:263-269', 'Eliminar el aporte de precio a confidence'],
 ['scanner.py:215-222', 'Degradar precio a penalización anti-late'],
 ['velocity.py:133', 'rising deja de ser obligatorio; añadir rama "vol sube, precio plano"'],
 ['main.py:1029-1034', 'Sustituir auto-enter inmediato por transición a WATCHLIST'],
 ['main.py lifespan:206', 'Añadir _observe_loop'],
 ['(nuevos)', 'microstructure.py, accumulation.py, persistence.py, sequence.py, rug_risk.py'],
], widths=[1.7, 4.5])

h3('5.5 Plan por etapas')
table(['Etapa', 'Qué', 'Salida / criterio'], [
 ['1', 'Datos (2-4 sem). microstructure.py + _observe_loop + tabla. Sin operar.', 'Series de candidatos Y de pumps confirmados, etiquetados'],
 ['2', 'Etiquetado + EDA. N≥200 eventos.', '¿La acumulación PRECEDE al pump aquí? El edge se confirma o se mata.'],
 ['3', 'Construir scores sobre datos REALES (no sim).', 'accumulation/persistence/sequence/rug calibrados'],
 ['4', 'Backtest real + ablation + edge-search (ahora SÍ válidos).', 'Universo A/B/C, ventana 15/30/60'],
 ['5', 'Paper forward-test ≥4 sem.', 'Solo si EV>0 robusto fuera de muestra'],
 ['6', 'Live opt-in, claves propias, tamaño mínimo.', 'Solo tras Etapa 5 positiva'],
], widths=[0.5, 3.0, 2.7])

h3('5.6 Riesgos por cambio')
table(['Cambio', 'Riesgo'], [
 ['Invertir universo', 'Explosión de laterales (la mayoría no pumpea) → ruido. Mitiga: persistence como filtro duro'],
 ['Quitar precio del score', 'Pierdes el único feature con señal PROBADA (aunque tardía). Cambiar −EV conocido por 0-EV desconocido'],
 ['Fase de observación', 'Latencia: si el move es rápido te lo pierdes. Trade-off central anticipación vs confirmación'],
 ['RugRiskScore temporal', 'Falsos positivos cortan ganadores legítimos'],
], widths=[1.7, 4.5])

h3('5.7 Impacto esperado (honesto)')
callout('NO se pueden dar números de EV/PF/WR/DD. Cualquier cifra antes de la Etapa 4 sería inventada — '
        'y inventar datos viola la regla del proyecto.', AMBER)
p('Dirección esperada SI la hipótesis de acumulación tiene señal: WR baja (entras antes, más '
  'incertidumbre), pero avg-win sube (compras antes del move) y la cola de hard-stop se reduce (sales '
  'del libro fino + rug-score). Eso PODRÍA voltear PF>1. Es hipótesis, no promesa.')

# ===== RESPUESTAS OBLIGATORIAS =====
h1('6. Respuestas obligatorias')
h3('6.1 ¿Qué detecta — A/B/C/D/E?')
callout('Respuesta: B + C + algo de D → una mezcla, dominada por B/C.', INK)
table(['Opción', '¿?', 'Justificación'], [
 ['A · preparación de pump', 'NO', 'Universo ciego a laterales; score exige precio movido'],
 ['B · inicio de pump', 'SÍ', 'Camino velocity dispara en la ignición'],
 ['C · pumps tardíos', 'SÍ', 'Universo top-gainer + anti-chase admite tokens ya +30-60%'],
 ['D · manipulación', 'marginal', 'Solo el flag estático top_book_share>0.80'],
 ['E · mezcla', 'SÍ', 'Es B+C dominante con algo de D'],
], widths=[1.9, 0.8, 3.5])
p('Justificación: triple gate de precio (universo change>0, score ∝chg, velocity rising) + el propio '
  'anti-chase confesando que llega tarde.', italic=True, color=GREY)

h3('6.2 ¿La hipótesis original tiene edge estadístico?')
p('Dos hipótesis distintas, dos respuestas:', bold=True)
callout('HIPÓTESIS ORIGINAL (detectar momentum de microcaps): sin evidencia de edge; evidencia '
        'direccional de que es estructuralmente NEGATIVA. No por mala implementación — el '
        'position_manager está bien hecho y aun así pierde. Pierde porque el trigger del bot y el '
        'trigger del rug son el MISMO evento. Es propiedad del fenómeno, no de los parámetros. '
        'Ninguna optimización de umbrales lo arregla.', RED)
callout('HIPÓTESIS NUEVA (detectar preparación/acumulación): NO falsada, pero tampoco validada — es '
        'DESCONOCIDA. Es la única rama que el sistema nunca probó, porque nunca grabó los datos para '
        'probarla. Podría tener edge; podría no tenerlo. Hoy no hay base para afirmar ninguna.', AMBER)

h3('6.3 Recomendación de Principal Quant')
p('No construir los 5 módulos a ciegas (sería el "parche elegante" que pediste evitar, solo que más '
  'caro). Construir ÚNICAMENTE la Etapa 1 (recolector), grabar 3-4 semanas, y dejar que los datos '
  'decidan si la acumulación precede al pump en MEXC/Bitget.', bold=True)
bullet('Si los datos dicen que sí → tienes un sistema con base real.', None)
bullet('Si dicen que no → te ahorraste construir 5 módulos sobre una hipótesis falsa.', None)
bullet('Mantener PAPER mientras tanto. Live solo tras Etapa 5 positiva.', None)

doc.add_paragraph()
p('— Fin del informe. Auditoría sobre código fuente real de apps/pump-reader. Sin datos de negocio '
  'inventados. Donde un dato no existe, se dice explícitamente.', italic=True, color=GREY, size=9)

out = r'C:\Users\osval\OneDrive\Escritorio\TradeOS_Auditoria_Estructural_Rediseno.docx'
doc.save(out)
print('SAVED', out)
print('paragraphs', len(doc.paragraphs), 'tables', len(doc.tables))
