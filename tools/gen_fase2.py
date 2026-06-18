# -*- coding: utf-8 -*-
"""Documento de especificacion tecnica FASE 2 (diseno, SIN codigo).
AccumulationScore / PersistenceScore / RugRiskScore / Sequence Engine /
maquina de estados Candidate->Watchlist->Monitor->Entry. + hallazgos de Fase 1."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INK = RGBColor(0x0c,0x10,0x18); PINK = RGBColor(0xff,0x2f,0x6e); GREY = RGBColor(0x55,0x5b,0x66)
GREEN = RGBColor(0x18,0x8a,0x5a); RED = RGBColor(0xc0,0x2a,0x2a); AMBER = RGBColor(0xb5,0x6a,0x00)
doc = Document(); base = doc.styles['Normal']; base.font.name='Calibri'; base.font.size=Pt(10.5)

def _shade(c,h):
    tcPr=c._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),h); tcPr.append(shd)
def h1(t):
    p=doc.add_heading(t,level=1)
    for r in p.runs: r.font.color.rgb=PINK
    return p
def h2(t):
    p=doc.add_heading(t,level=2)
    for r in p.runs: r.font.color.rgb=INK
    return p
def h3(t):
    p=doc.add_heading(t,level=3)
    for r in p.runs: r.font.color.rgb=GREY
    return p
def p(t,bold=False,italic=False,color=None,size=10.5):
    par=doc.add_paragraph(); r=par.add_run(t); r.bold=bold; r.italic=italic; r.font.size=Pt(size)
    if color: r.font.color.rgb=color
    return par
def callout(t,color=INK):
    par=doc.add_paragraph(); par.paragraph_format.left_indent=Inches(0.15)
    par.paragraph_format.space_before=Pt(4); par.paragraph_format.space_after=Pt(6)
    r=par.add_run(t); r.bold=True; r.font.size=Pt(10); r.font.color.rgb=color; return par
def bullet(t,prefix=None):
    par=doc.add_paragraph(style='List Bullet')
    if prefix:
        r=par.add_run(prefix); r.bold=True; par.add_run(t)
    else: par.add_run(t)
    return par
def formula(t):
    par=doc.add_paragraph(); par.paragraph_format.left_indent=Inches(0.2)
    par.paragraph_format.space_before=Pt(2); par.paragraph_format.space_after=Pt(6)
    for line in t.split('\n'):
        r=par.add_run(line+'\n'); r.font.name='Consolas'; r.font.size=Pt(9); r.font.color.rgb=RGBColor(0x1b,0x2a,0x4a)
    return par
def table(headers,rows,widths=None):
    tb=doc.add_table(rows=1,cols=len(headers)); tb.style='Light Grid Accent 1'; tb.alignment=WD_TABLE_ALIGNMENT.LEFT
    hc=tb.rows[0].cells
    for i,ht in enumerate(headers):
        hc[i].text=''; rr=hc[i].paragraphs[0].add_run(ht); rr.bold=True; rr.font.size=Pt(9)
        rr.font.color.rgb=RGBColor(0xff,0xff,0xff); _shade(hc[i],'2f3a52')
    for row in rows:
        cells=tb.add_row().cells
        for i,val in enumerate(row):
            cells[i].text=''; rr=cells[i].paragraphs[0].add_run(str(val)); rr.font.size=Pt(8.7)
    if widths:
        for i,w in enumerate(widths):
            for row in tb.rows: row.cells[i].width=Inches(w)
    doc.add_paragraph(); return tb

# ===== PORTADA =====
t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=t.add_run('Amatista · TradeOS'); r.bold=True; r.font.size=Pt(28); r.font.color.rgb=PINK
s=doc.add_paragraph(); s.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=s.add_run('Especificación Técnica — FASE 2'); r.font.size=Pt(15); r.font.color.rgb=INK
s2=doc.add_paragraph(); s2.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=s2.add_run('AccumulationScore · PersistenceScore · RugRiskScore · Sequence Engine · Máquina de estados')
r.italic=True; r.font.size=Pt(10.5); r.font.color.rgb=GREY
m=doc.add_paragraph(); m.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=m.add_run('DISEÑO solamente — sin código. Se implementa tras validar con los datos de Fase 1. 2026-06-18')
r.font.size=Pt(9); r.font.color.rgb=GREY
p('')
callout('REGLA: nada de esta Fase 2 se implementa hasta tener una base de datos de observación limpia '
        '(Fase 1) y haber confirmado, con datos reales, que la acumulación precede al pump. Esto es la '
        'especificación que se calibrará y validará contra esos datos — no una promesa de edge.', AMBER)
doc.add_page_break()

# ===== 1 =====
h1('1. Propósito y encaje en la arquitectura')
p('Fase 2 transforma el sistema de "detector de momentum" a "detector de preparación". Consume '
  'EXCLUSIVAMENTE la tabla micro_snapshots que graba Fase 1 (no añade fuentes de datos). Cada score '
  'opera sobre una VENTANA temporal de un símbolo, no sobre un snapshot.')
callout('Contrato de entrada de todos los módulos: una serie ordenada de MicroSnapshots de un mismo '
        'símbolo/exchange, W = [s(t-k), …, s(t)], con k configurable (15/30/60 min). Los scores son '
        'funciones puras de W → entero 0-100.', INK)
p('Esto resuelve el defecto estructural #3 de la auditoría (el sistema evalúa fotos, no películas).')

# ===== 2 =====
h1('2. AccumulationScore — detectar acumulación silenciosa')
p('Objetivo: detectar que alguien está COMPRANDO sin mover el precio (la firma de la preparación). '
  'NO detecta pumps. Variables (todas en micro_snapshots):')
table(['Componente', 'Señal medida', 'Peso prop.', 'Racional'], [
 ['Vol↑ + precio plano', 'pendiente(volume) > 0 con |Δprecio ventana| < ~2%', '30', 'La firma central: volumen entra, precio no se mueve = absorción'],
 ['Profundidad bid↑', 'pendiente(bid_depth) > 0', '20', 'Alguien apila órdenes de compra progresivamente'],
 ['Absorción de ventas', 'precio baja intrabar pero el bid repone (ratio 0..1)', '20', 'Las ventas se comen sin que el precio ceda'],
 ['Spread comprimiéndose', 'pendiente(spread_pct) < 0', '15', 'Un market maker toma control del libro'],
 ['Imbalance comprador sano', '0,60 ≤ media(imbalance) ≤ 0,85 (sin blow-off)', '15', 'Sesgo comprador sostenido, no euforia terminal'],
], widths=[1.7, 2.4, 0.8, 2.0])
h3('Forma funcional (spec, no implementación)')
formula('AccumulationScore = clamp( Σ wᵢ · fᵢ(W) , 0, 100 )\n'
        '   donde fᵢ ∈ [0,1] (señal normalizada) y Σwᵢ = 100\n'
        '   pesos wᵢ = {30, 20, 20, 15, 15}  (a calibrar con Fase 1)')
p('Normalización: cada fᵢ se mapea a [0,1] con umbrales derivados de la distribución empírica de '
  'micro_snapshots (percentiles), no con constantes inventadas. La calibración es parte de la validación.',
  italic=True, color=GREY)

# ===== 3 =====
h1('3. PersistenceScore — la señal sostenida en el tiempo')
p('Objetivo: distinguir una condición que se MANTIENE de un parpadeo de un solo intervalo. Un imbalance '
  '0,7 durante 20 minutos vale mucho más que 0,7 en un snapshot.')
table(['Componente', 'Definición', 'Peso prop.'], [
 ['Imbalance sostenido', 'fracción de intervalos de W con imbalance > 0,65', '25'],
 ['Volumen creciente', 'grado de monotonía creciente de volume sobre W', '25'],
 ['Liquidez creciente', 'grado de monotonía creciente de liquidity_usd sobre W', '25'],
 ['Absorción sostenida', 'fracción de intervalos con absorción positiva', '25'],
], widths=[1.9, 3.6, 0.8])
formula('PersistenceScore = clamp( 25·g_imb + 25·g_vol + 25·g_liq + 25·g_abs , 0, 100 )\n'
        '   gᵢ ∈ [0,1]; "monotonía" = nº de pasos crecientes / (|W|-1)')
callout('PersistenceScore es el filtro anti-ruido del rediseño: sin él, ampliar el universo a tokens '
        'laterales (Hallazgo #1) ahogaría la señal en falsos positivos.', INK)

# ===== 4 =====
h1('4. RugRiskScore — probabilidad de colapso (modelo independiente)')
p('Regla dura: NUNCA un score único. PumpScore (alza) y RugRiskScore (colapso) son dos ejes ortogonales. '
  'RugRiskScore usa DETERIORO TEMPORAL — imposible hoy, posible con la serie de Fase 1.')
table(['Señal de riesgo', 'Medición sobre W', 'Peso prop.', 'Por qué'], [
 ['Liquidez desapareciendo', 'pendiente(liquidity_usd) < 0', '25', 'El libro se vacía antes del rug'],
 ['Spread ensanchando', 'pendiente(spread_pct) > 0', '20', 'Los market makers se retiran'],
 ['Concentración alta', 'media(top_book_share) > 0,8', '20', 'Libro sostenido por pocas órdenes (actor único)'],
 ['Movimiento sin soporte', 'velas con |Δprecio| alto y profundidad baja detrás', '25', 'Precio empujado sin libro real = trampa'],
 ['Retirada súbita de bid', 'caída abrupta de bid_depth entre intervalos', '10', 'Spoof/exit del manipulador'],
], widths=[1.8, 2.6, 0.8, 1.5])
formula('RugRiskScore = clamp( Σ vⱼ · hⱼ(W) , 0, 100 )\n'
        'Decisión combinada (matriz, NO suma):\n'
        '   ENTRAR  sii  PumpScore ≥ P   AND   RugRiskScore ≤ R\n'
        '   (P, R) se eligen del backtest de Fase 1 — frontera de Pareto EV vs drawdown')
p('Esto reemplaza el actual forensic_check (gate binario estático) por un modelo continuo y temporal. '
  'El forensic_check de hoy queda como caso degenerado (snapshot, sin pendientes).', italic=True, color=GREY)

# ===== 5 =====
h1('5. Sequence Engine — la FORMA importa más que el nivel')
p('Premisa: una rampa 1,5x → 2x → 3x → 5x vale más que "5x ahora mismo". El Sequence Engine puntúa la '
  'TRAYECTORIA, no el último valor.')
table(['Entrada', 'Qué evalúa', 'Salida'], [
 ['Serie de volume (o de cualquier métrica)', 'Ratios sucesivos rₜ = xₜ / xₜ₋₁', 'sequence_bonus ∈ [0,1]'],
 ['', '¿Los ratios son crecientes (aceleración real)?', 'multiplicador 1,0 vs 0,3'],
 ['', '¿Magnitud media de la rampa?', 'clamp(media(r)-1, 0, 1)'],
], widths=[2.4, 2.8, 1.0])
formula('sequence_bonus = clamp( media(rₜ) − 1 , 0, 1 ) × ( 1 si rₜ creciente, si no 0,3 )\n'
        'Se aplica como multiplicador/bonus a Accumulation y Persistence, no como score aparte.')
p('El Sequence Engine es transversal: convierte cualquier métrica de la serie (volumen, liquidez, '
  'imbalance) en una evaluación de tendencia con memoria. Es la pieza que hace que el sistema "vea '
  'películas".')

# ===== 6 =====
h1('6. Máquina de estados: Candidate → Watchlist → Monitor → Entry')
p('Prohíbe la entrada inmediata (defecto estructural actual: el scan entra en la misma pasada que '
  'detecta). Todo candidato se OBSERVA antes de poder entrar.')
table(['Estado', 'Disparador de entrada al estado', 'Qué se evalúa', 'Transición'], [
 ['CANDIDATE', 'Aparece en el scan / universo ampliado', 'Filtros mínimos de universo', '→ WATCHLIST'],
 ['WATCHLIST', 'Admitido a observación', 'Se graba su micro_snapshot cada minuto (Fase 1)', '→ MONITOR tras N min'],
 ['MONITOR', 'Hay ventana suficiente (15/30/60 min)', 'AccumulationScore, PersistenceScore, RugRiskScore sobre W', '→ CONFIRMATION o descarte'],
 ['CONFIRMATION', 'Scores cruzan umbral Y se sostienen', 'Que la señal persista K intervalos sin blow-off ni rug', '→ ENTRY o expira'],
 ['ENTRY', 'Confirmado', 'Entrega al motor de ejecución actual (sin cambios)', '(posición gestionada por PositionManager actual)'],
], widths=[1.2, 1.9, 2.0, 1.4])
callout('Decisión de ventana (15/30/60 min): NO se fija en el diseño. Se elige con los datos de Fase 1 '
        'midiendo qué ventana da más lead-time con suficiente precisión. Es un resultado empírico, no '
        'una constante de diseño.', AMBER)
p('La máquina de estados NO reemplaza el PositionManager ni el RiskGuard: una vez en ENTRY, la gestión '
  'de la posición (TP/SL/trailing/hard-stop) es la actual, que la auditoría calificó como correcta.')

# ===== 7 =====
h1('7. Plan de validación (cómo se calibra con Fase 1)')
table(['Paso', 'Acción', 'Criterio de avance'], [
 ['V1', 'Etiquetar pumps reales en micro_snapshots (MFE ≥ X% en H horas)', '≥ N eventos etiquetados (objetivo N≥200)'],
 ['V2', 'Para cada pump, reconstruir W de -180..0 min', 'Cobertura: ¿hay pre-historia? (ver Hallazgo #1)'],
 ['V3', 'Medir si Accumulation/Persistence suben ANTES del pump vs control', 'Separación estadística significativa o se DESCARTA'],
 ['V4', 'Calibrar pesos/umbrales por percentiles empíricos', 'Sin overfitting: train/test temporal separado'],
 ['V5', 'Backtest con (P,R) + ventana → EV/PF/lead-time fuera de muestra', 'EV>0 robusto o no se promociona a paper'],
], widths=[0.6, 3.3, 2.4])
callout('Si V3 falla (la acumulación NO precede al pump en estos exchanges), Fase 2 NO se construye. Ese '
        'es el valor de hacer Fase 1 primero: poder matar la hipótesis barato.', RED)

# ===== 8 =====
h1('8. Hallazgos arquitectónicos detectados en Fase 1 (documentar, NO corregir)')
p('Por la regla acordada: se registran, no se tocan, hasta tener la base de datos limpia.', italic=True, color=GREY)
table(['#', 'Hallazgo', 'Impacto', 'Acción diferida'], [
 ['1', 'La watchlist de observación se alimenta SOLO de candidatos del scan, que es gainer-gated (change>0). Un token verdaderamente lateral no se observa hasta que YA empezó a moverse.', 'Sesgo de cobertura: la pre-historia de pumps que nacen de quietud total puede faltar o llegar tarde. Es el mismo chicken-and-egg de la auditoría.', 'Fase 1.5 (opcional): añadir un universo de observación más amplio (muestra rotatoria de tokens de volumen medio y precio lateral), independiente del scan. NO implementar aún.'],
 ['2', 'El scan hace _candidates.clear() cada pasada y el snapshot se sobrescribe: el bot nunca tuvo memoria temporal propia.', 'Fase 1 lo soluciona con DB externa, pero el scan sigue sin estado. Cualquier lógica futura debe leer la DB, no _candidates.', 'Diseño de Fase 2 ya asume lectura desde micro_snapshots, no desde _candidates.'],
 ['3', 'velocity calcula accel pero no lo persiste; Fase 1 recomputa "velocity" desde su propio buffer.', 'Posible discrepancia menor entre el accel del trigger y el velocity grabado (distintos baselines).', 'Documentado. Unificar definición de velocity en Fase 2 si se confirma que aporta.'],
 ['4', 'No hay etiqueta de "pump confirmado" persistida en una tabla; LearningLab la tiene en memoria y se pierde al reiniciar.', 'V1 de validación necesita etiquetas estables. Habrá que persistir outcomes.', 'Fase 1.5: persistir outcomes de LearningLab a la DB local (tabla aparte). NO implementar aún.'],
], widths=[0.3, 2.6, 2.0, 1.9])

doc.add_paragraph()
p('— Fin de la especificación FASE 2. Diseño solamente. Implementación condicionada a la validación '
  'empírica con los datos de Fase 1. Sin datos de negocio inventados.', italic=True, color=GREY, size=9)

out=r'C:\Users\osval\OneDrive\Escritorio\TradeOS_Fase2_Especificacion.docx'
doc.save(out); print('SAVED',out); print('paragraphs',len(doc.paragraphs),'tables',len(doc.tables))
